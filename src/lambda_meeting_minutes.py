"""
Lambda: Meeting Minutes Generator v1.1

Changes from v1.0:
- REFACTOR: Transcript parsing delegated to shared transcript_utils.py
- FIX: Unified time extraction — handles both VAD segments and full audio correctly
- FIX: Per-speaker-turn absolute timestamps derived from Transcribe word items
- CHANGE: Default CLAUDE_MODEL to claude-sonnet-4-6

Generic meeting minutes generation from transcribed audio recordings.
Not site-inspection-specific — handles any business meeting context:
brainstorming, strategy, BD, standups, retrospectives, etc.

This Lambda function:
1. Collects all transcripts from target meeting date/session
2. Optionally loads meeting config (title, attendees, type) from S3 or event payload
3. Calls Claude API to generate STRUCTURED meeting minutes JSON
4. Saves debug record (prompt + response) for prompt tuning
5. Writes minutes JSON + Word doc to S3

Trigger:
  - API Gateway (on-demand)  → {"date": "2026-03-20"}
  - API Gateway (with config) → {"date": "2026-03-20", "meeting_title": "...", "attendees": [...]}
  - EventBridge (scheduled)   → {"date": "yesterday"}

Event payload options:
  date              — Target date (YYYY-MM-DD) or "yesterday" or "today"
  meeting_title     — Optional: override meeting title
  meeting_type      — Optional: strategy | standup | brainstorm | review | bd | general
  attendees         — Optional: list of attendee names
  user              — Optional: specific device/user folder to process (default: all)
  transcript_prefix — Optional: custom S3 prefix for transcripts (e.g. "meetings/2026-03-20/")

Environment Variables:
    S3_BUCKET           - S3 bucket name
    ANTHROPIC_API_KEY   - Anthropic API key (sk-ant-xxx)
    CLAUDE_MODEL        - Claude model ID (default: claude-sonnet-4-5-20250929)
    MINUTES_PREFIX      - Output prefix (default: meeting_minutes/)
"""

import os
import json
import logging
import re
import boto3
import urllib3
from datetime import datetime, timedelta
from io import BytesIO
from transcript_utils import (
    normalize_transcript, format_turns_for_prompt, get_time_bounds,
    extract_device_from_filename, write_meeting_manifest,
)

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
s3_client = boto3.client('s3')

# Configuration
S3_BUCKET = os.environ.get('S3_BUCKET', '')
MINUTES_PREFIX = os.environ.get('MINUTES_PREFIX', 'meeting_minutes/')
REPORT_PREFIX = os.environ.get('REPORT_PREFIX', 'reports/')  # compat output path
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
CLAUDE_MODEL = os.environ.get('CLAUDE_MODEL', 'claude-sonnet-4-6')

# Check python-docx availability once at import time
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
    logger.info("python-docx available — Word generation enabled")
except ImportError:
    logger.warning(
        "python-docx NOT installed — Word generation disabled. "
        "To enable: add a Lambda Layer with python-docx."
    )


# ============================================================
# JSON Schemas
# ============================================================

MEETING_MINUTES_SCHEMA = """{
  "executive_summary": [
    "Bullet 1: Meeting purpose and context",
    "Bullet 2: Key outcome or decision",
    "Bullet 3: Overall direction agreed or next major milestone"
  ],
  "topics": [
    {
      "topic_id": 0,
      "time_range": "HH:MM – HH:MM",
      "topic_title": "Short descriptive title",
      "category": "strategy | operations | finance | product | partnership | technical | hr | legal | general",
      "status": "decided | deferred | in_discussion | blocked",
      "participants": ["Person A", "Person B"],
      "summary": "2-4 sentence summary of the discussion on this topic",
      "key_decisions": [
        {
          "decision": "What was decided",
          "rationale": "Why (1 sentence, optional — omit key if not stated)",
          "decided_by": "Person or group"
        }
      ],
      "action_items": [
        {
          "action": "What needs to be done",
          "owner": "Person name",
          "deadline": "When (e.g. 'By Friday', 'EOW', '2026-03-25', 'ASAP')",
          "priority": "high | medium | low"
        }
      ],
      "open_questions": [
        "Question or unresolved point that needs follow-up"
      ]
    }
  ],
  "follow_ups": [
    {
      "item": "What needs to happen next",
      "owner": "Person name",
      "deadline": "When",
      "depends_on": "Blocking dependency (optional — omit key if none)",
      "priority": "high | medium | low"
    }
  ],
  "next_steps": [
    "Top 3-5 most important actions after this meeting"
  ],
  "parking_lot": [
    "Items raised but explicitly deferred to a future meeting"
  ]
}"""

WEEKLY_MEETINGS_SCHEMA = """{
  "executive_summary": "3-5 sentence overview of the week's meetings",
  "meetings_covered": [
    {
      "date": "YYYY-MM-DD",
      "title": "Meeting title",
      "key_outcome": "One sentence"
    }
  ],
  "key_decisions": [
    {
      "decision": "What was decided",
      "meeting_date": "YYYY-MM-DD",
      "status": "final | provisional | reversed"
    }
  ],
  "outstanding_actions": [
    {
      "action": "What needs to be done",
      "owner": "Person name",
      "original_date": "YYYY-MM-DD",
      "deadline": "When",
      "priority": "high | medium | low",
      "status": "open | overdue | completed"
    }
  ],
  "recurring_themes": ["Theme 1", "Theme 2"],
  "next_week_priorities": ["Priority 1", "Priority 2"]
}"""


# ============================================================
# Date Helpers
# ============================================================

def get_nzdt_now():
    """Get current time in NZDT (UTC+13)"""
    return datetime.utcnow() + timedelta(hours=13)


def resolve_date(date_str):
    """Resolve 'today', 'yesterday', or YYYY-MM-DD to a date string."""
    if not date_str or date_str == 'yesterday':
        return (get_nzdt_now() - timedelta(days=1)).strftime('%Y-%m-%d')
    if date_str == 'today':
        return get_nzdt_now().strftime('%Y-%m-%d')
    return date_str


# ============================================================
# S3 Helpers
# ============================================================

def list_s3_objects(bucket, prefix):
    """List all objects under a prefix"""
    objects = []
    paginator = s3_client.get_paginator('list_objects_v2')
    for page in paginator.paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get('Contents', []):
            objects.append({
                'key': obj['Key'],
                'last_modified': obj['LastModified'],
                'size': obj['Size']
            })
    return objects


def download_json_from_s3(bucket, key):
    """Download and parse JSON file from S3"""
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read().decode('utf-8')
        return json.loads(content)
    except Exception as e:
        logger.error(f"Failed to download {key}: {str(e)}")
        return None


# ============================================================
# User / Attendee Mapping
# ============================================================

_user_mapping_cache = None


def load_user_mapping(bucket):
    """
    Load user mapping from S3 config/user_mapping.json.
    Returns: {"device_id": "Display Name", ...}
    """
    global _user_mapping_cache
    if _user_mapping_cache is not None:
        return _user_mapping_cache

    try:
        data = download_json_from_s3(bucket, 'config/user_mapping.json')
        if data:
            raw_mapping = data.get('mapping', {})
            normalized = {}
            for device, value in raw_mapping.items():
                if isinstance(value, str):
                    normalized[device] = value
                elif isinstance(value, dict):
                    normalized[device] = value.get('name', device)
                else:
                    normalized[device] = str(value)
            _user_mapping_cache = normalized
            logger.info(f"Loaded user mapping: {len(normalized)} entries")
            return _user_mapping_cache
    except Exception as e:
        logger.warning(f"Failed to load user mapping: {e}")

    _user_mapping_cache = {}
    return _user_mapping_cache


# ============================================================
# Transcript Parsing — delegated to transcript_utils.py
# Functions used: normalize_transcript(), format_turns_for_prompt(),
#                 get_time_bounds(), extract_device_from_filename()
# ============================================================


# ============================================================
# Collect Transcripts for a Date
# ============================================================

def collect_transcripts(bucket, target_date, user_filter=None, custom_prefix=None):
    """
    Collect all transcripts for a given date, normalized via transcript_utils.

    Each entry in the returned list is the output of normalize_transcript(),
    with an additional 'key' field for the S3 object path.

    Returns: list of normalized transcript dicts sorted by segment_base_time
    """
    user_mapping = load_user_mapping(bucket)
    transcripts = []

    def _process_object(key):
        """Download, normalize, and append a single transcript JSON."""
        filename = os.path.basename(key)
        data = download_json_from_s3(bucket, key)
        if not data:
            return
        normalized = normalize_transcript(data, filename, user_mapping=user_mapping)
        if normalized and normalized.get('full_text'):
            normalized['key'] = key
            transcripts.append(normalized)

    if custom_prefix:
        for obj in list_s3_objects(bucket, custom_prefix):
            if obj['key'].endswith('.json'):
                _process_object(obj['key'])
    else:
        # Discover users from transcripts/ folder
        users = set()
        if user_filter:
            users.add(user_filter)
        else:
            for obj in list_s3_objects(bucket, 'transcripts/'):
                parts = obj['key'].split('/')
                if len(parts) >= 2 and parts[1]:
                    users.add(parts[1])

        for user_name in sorted(users):
            # Date subfolder first
            prefix = f"transcripts/{user_name}/{target_date}/"
            transcript_objects = list_s3_objects(bucket, prefix)

            # Fallback: scan all and filter by date
            if not transcript_objects:
                prefix_all = f"transcripts/{user_name}/"
                transcript_objects = [
                    obj for obj in list_s3_objects(bucket, prefix_all)
                    if target_date in obj['key']
                ]

            for obj in transcript_objects:
                key = obj['key']
                if not key.endswith('.json'):
                    continue
                if target_date not in key:
                    continue
                _process_object(key)

    transcripts.sort(key=lambda x: x.get('segment_base_time') or datetime.min)
    logger.info(f"Collected {len(transcripts)} transcripts for {target_date}")
    return transcripts


# ============================================================
# Prompt Builder
# ============================================================

def build_meeting_prompt(transcripts, meeting_config):
    """
    Build the Claude prompt for meeting minutes.

    Args:
        transcripts: sorted list of transcript entries
        meeting_config: dict with title, date, attendees, meeting_type, etc.
    """
    target_date = meeting_config.get('date', '?')
    meeting_title = meeting_config.get('meeting_title', f'Meeting — {target_date}')
    meeting_type = meeting_config.get('meeting_type', 'general')
    attendees = meeting_config.get('attendees', [])
    explicit_attendees = bool(attendees)  # True if user provided attendee list

    # Load prompt template if available
    prompt_template = None
    try:
        tmpl_data = download_json_from_s3(S3_BUCKET, 'config/prompt_templates_meeting.json')
        if tmpl_data:
            prompt_template = tmpl_data.get('meeting_minutes', {})
            logger.info("Loaded meeting prompt template from S3")
    except Exception:
        pass

    # Build metadata block
    metadata_lines = [
        f"Date: {target_date}",
        f"Meeting Title: {meeting_title}",
        f"Meeting Type: {meeting_type}",
    ]
    if attendees:
        metadata_lines.append(f"Attendees: {', '.join(attendees)}")

    # Duration from actual per-turn timestamps via transcript_utils
    earliest, latest, duration_mins = get_time_bounds(transcripts)
    if earliest and latest:
        metadata_lines.append(f"Duration: ~{duration_mins} minutes")
        metadata_lines.append(
            f"Recording Window: {earliest.strftime('%H:%M')} – {latest.strftime('%H:%M')}"
        )

    total_words = sum(t.get('word_count', 0) for t in transcripts)
    metadata_lines.append(f"Total Recordings: {len(transcripts)}")
    metadata_lines.append(f"Total Words: {total_words}")

    metadata_block = '\n'.join(metadata_lines)

    # Build attendee/speaker reference
    speakers = {}
    for t in transcripts:
        device = t.get('device', 'Unknown')
        name = t.get('speaker_name', device)
        if device not in speakers:
            speakers[device] = name

    ref_lines = []
    if explicit_attendees:
        ref_lines.append("Recording Devices:")
        for device, mapped_name in speakers.items():
            ref_lines.append(f"  {device} (device owner: {mapped_name}) — recording device")
        ref_lines.append("")
        ref_lines.append(f"IMPORTANT: The only people in this meeting are: {', '.join(attendees)}")
        ref_lines.append("Speaker diarization labels (spk_0, spk_1, etc.) correspond to these attendees.")
        ref_lines.append("Use ONLY names from the Attendees list for participants — NOT device owner names.")
    elif speakers:
        ref_lines.append("Speaker Reference:")
        for device, name in speakers.items():
            ref_lines.append(f"  {device} = {name}")

    attendee_reference = '\n'.join(ref_lines)

    # Build transcript text via transcript_utils — per-turn absolute timestamps
    transcript_lines = []
    for t in transcripts:
        label = t.get('device', 'Unknown') if explicit_attendees else None
        lines = format_turns_for_prompt(t, label_override=label, use_absolute_time=True)
        transcript_lines.extend(lines)

    transcripts_text = '\n\n'.join(transcript_lines)

    # Use S3 prompt template if loaded, otherwise use inline default
    if prompt_template and prompt_template.get('prompt'):
        prompt_body = prompt_template['prompt']
        prompt_body = prompt_body.replace('{metadata_block}', metadata_block)
        prompt_body = prompt_body.replace('{attendee_reference}', attendee_reference)
        prompt_body = prompt_body.replace('{transcripts_text}', transcripts_text[:120000])
        prompt_body = prompt_body.replace('{schema}', MEETING_MINUTES_SCHEMA)

        system_context = prompt_template.get('system_context', '')
        return f"{system_context}\n\n{prompt_body}" if system_context else prompt_body

    # Inline fallback prompt
    return f"""You are a professional meeting minutes assistant. You produce clear, actionable meeting records from audio transcripts.

## Meeting Info
{metadata_block}
{attendee_reference}

Analyze the following meeting transcript and produce STRUCTURED meeting minutes.

## Transcript (chronological)
{transcripts_text[:120000]}

## Instructions
1. Write an executive_summary as a bullet-point array (3-5 bullets). Cover: meeting purpose, key outcomes, overall direction.
2. Identify and group the transcript into logical AGENDA TOPICS. If no formal agenda was followed, infer topics from conversation flow.
3. For each topic, extract:
   - A concise summary of the discussion
   - Specific decisions made (verbatim intent, not paraphrased away)
   - Action items with owner, deadline, and priority
   - Open questions or unresolved points that need follow-up
4. Classify each topic into one of: strategy, operations, finance, product, partnership, technical, hr, legal, general
5. Identify participants for each topic — use ONLY names from the Attendees list if provided. Map speaker diarization labels to attendees based on conversation context.
6. Derive time_range for each topic from the segment timestamps in the transcript — do NOT just repeat the same start time.
7. If any topic has time-sensitive or blocking dependencies, flag them in the follow_ups array.
8. Compile a top-level next_steps array: the 3-5 most important things that must happen after this meeting.

## Output Format
Return ONLY valid JSON matching this exact schema (no markdown, no explanation, just the JSON object):

{MEETING_MINUTES_SCHEMA}

Rules:
- topic_id must be sequential starting from 0
- time_range format: "HH:MM – HH:MM" (use en-dash) — derive from the [HH:MM:SS – HH:MM:SS] timestamps shown in each transcript segment
- category MUST be one of: strategy, operations, finance, product, partnership, technical, hr, legal, general
- priority MUST be one of: high, medium, low
- status MUST be one of: decided, deferred, in_discussion, blocked
- executive_summary MUST be an array of bullet strings, NOT a single string
- action_items, open_questions, and follow_ups can be empty arrays []
- participants MUST only contain names from the Attendees list — never use device IDs or device owner names
- Do NOT include any text outside the JSON object
- Do NOT fabricate decisions or action items not clearly stated in the transcript
- Preserve the original language of decisions — do not soften or editorialize"""


# ============================================================
# Claude API
# ============================================================

def call_claude_structured(prompt, max_tokens=4096):
    """Call Anthropic Claude API. Returns (raw_text, error_string)."""
    if not ANTHROPIC_API_KEY:
        logger.error("ANTHROPIC_API_KEY not set")
        return None, "ANTHROPIC_API_KEY not configured"

    http = urllib3.PoolManager()
    body = json.dumps({
        "model": CLAUDE_MODEL,
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}]
    })

    try:
        resp = http.request(
            'POST', 'https://api.anthropic.com/v1/messages',
            body=body,
            headers={
                'Content-Type': 'application/json',
                'x-api-key': ANTHROPIC_API_KEY,
                'anthropic-version': '2023-06-01',
            },
            timeout=180.0,
        )
        data = json.loads(resp.data.decode('utf-8'))

        if resp.status == 200:
            text_blocks = [b['text'] for b in data.get('content', []) if b.get('type') == 'text']
            raw_text = '\n'.join(text_blocks)
            return raw_text, None
        else:
            err = data.get('error', {}).get('message', f'HTTP {resp.status}')
            logger.error(f"Claude API error: {err}")
            return None, err

    except Exception as e:
        logger.error(f"Claude API call failed: {str(e)}")
        return None, str(e)


def extract_json_from_response(raw_text):
    """Extract JSON from Claude's response."""
    # Try markdown fence
    match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw_text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # Try entire response
    try:
        return json.loads(raw_text.strip())
    except json.JSONDecodeError:
        pass

    # Try first { to last }
    first_brace = raw_text.find('{')
    last_brace = raw_text.rfind('}')
    if first_brace != -1 and last_brace != -1:
        try:
            return json.loads(raw_text[first_brace:last_brace + 1])
        except json.JSONDecodeError:
            pass

    logger.error(f"Failed to extract JSON from Claude response: {raw_text[:500]}")
    return None


# ============================================================
# Debug Record
# ============================================================

def save_debug_record(bucket, target_date, meeting_title, prompt, raw_response,
                      parsed_json, parse_success, input_stats):
    """Save debug record for prompt tuning."""
    try:
        safe_title = re.sub(r'[^a-zA-Z0-9_-]', '_', meeting_title)[:50]
        debug_key = f"{MINUTES_PREFIX}{target_date}/{safe_title}_debug.json"

        debug_record = {
            '_description': 'Debug record for meeting minutes prompt tuning.',
            'timestamp': datetime.utcnow().isoformat() + 'Z',
            'model': CLAUDE_MODEL,
            'meeting_title': meeting_title,
            'target_date': target_date,
            'parse_success': parse_success,
            'input_stats': input_stats,
            'prompt': prompt,
            'prompt_length': len(prompt),
            'raw_response': raw_response,
            'raw_response_length': len(raw_response) if raw_response else 0,
            'parsed_json': parsed_json,
        }

        s3_client.put_object(
            Bucket=bucket, Key=debug_key,
            Body=json.dumps(debug_record, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json'
        )
        logger.info(f"  Saved debug: {debug_key}")

    except Exception as e:
        logger.error(f"  Debug save failed: {e}")


# ============================================================
# Word Document Generation
# ============================================================

def generate_word_document(minutes_data, title):
    """Generate Word document from structured meeting minutes."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meeting info
    meeting_date = minutes_data.get('meeting_date', '')
    meeting_type = minutes_data.get('meeting_type', '')
    attendees = minutes_data.get('attendees', [])
    if meeting_date or meeting_type or attendees:
        info_parts = []
        if meeting_date:
            info_parts.append(f"Date: {meeting_date}")
        if meeting_type:
            info_parts.append(f"Type: {meeting_type}")
        if attendees:
            info_parts.append(f"Attendees: {', '.join(attendees)}")
        p = doc.add_paragraph(' | '.join(info_parts))
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = minutes_data.get('executive_summary', 'No summary available')
    if isinstance(exec_summary, list):
        for item in exec_summary:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph(exec_summary)

    # Topics
    topics = minutes_data.get('topics', [])
    if topics:
        doc.add_heading('Discussion Topics', level=1)

        for topic in topics:
            cat = topic.get('category', 'general').upper()
            status = topic.get('status', '').upper()
            time_range = topic.get('time_range', '')
            participants = topic.get('participants', [])

            header = f"{time_range}  {topic.get('topic_title', '')}  [{cat}]"
            if status:
                header += f"  — {status}"
            doc.add_heading(header, level=2)

            if participants:
                p = doc.add_paragraph(f"Participants: {', '.join(participants)}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

            # Summary
            doc.add_paragraph(topic.get('summary', ''))

            # Key Decisions
            decisions = topic.get('key_decisions', [])
            if decisions:
                doc.add_heading('Decisions', level=3)
                for d in decisions:
                    if isinstance(d, str):
                        doc.add_paragraph(d, style='List Bullet')
                    elif isinstance(d, dict):
                        text = d.get('decision', '')
                        rationale = d.get('rationale', '')
                        decided_by = d.get('decided_by', '')
                        line = text
                        if rationale:
                            line += f" — Rationale: {rationale}"
                        if decided_by:
                            line += f" (by {decided_by})"
                        doc.add_paragraph(line, style='List Bullet')

            # Action Items
            actions = topic.get('action_items', [])
            if actions:
                doc.add_heading('Action Items', level=3)
                for ai in actions:
                    priority = ai.get('priority', 'medium').upper()
                    text = (f"[{priority}] {ai.get('action', '')} "
                            f"→ {ai.get('owner', '?')} by {ai.get('deadline', '?')}")
                    p = doc.add_paragraph(text, style='List Bullet')
                    if priority == 'HIGH' and p.runs:
                        p.runs[0].font.color.rgb = RGBColor(192, 57, 43)

            # Open Questions
            questions = topic.get('open_questions', [])
            if questions:
                doc.add_heading('Open Questions', level=3)
                for q in questions:
                    doc.add_paragraph(f"❓ {q}", style='List Bullet')

    # Follow-ups
    follow_ups = minutes_data.get('follow_ups', [])
    if follow_ups:
        doc.add_heading('Follow-ups & Dependencies', level=1)
        for fu in follow_ups:
            priority = fu.get('priority', 'medium').upper()
            text = f"[{priority}] {fu.get('item', '')} → {fu.get('owner', '?')} by {fu.get('deadline', '?')}"
            depends = fu.get('depends_on', '')
            if depends:
                text += f"\n  ⚠ Blocked by: {depends}"
            doc.add_paragraph(text, style='List Bullet')

    # Next Steps
    next_steps = minutes_data.get('next_steps', [])
    if next_steps:
        doc.add_heading('Next Steps', level=1)
        for i, step in enumerate(next_steps, 1):
            doc.add_paragraph(f"{i}. {step}")

    # Parking Lot
    parking = minutes_data.get('parking_lot', [])
    if parking:
        doc.add_heading('Parking Lot (Deferred)', level=1)
        for item in parking:
            doc.add_paragraph(f"• {item}")

    # Metadata footer
    meta = minutes_data.get('_report_metadata', {})
    if meta:
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(
            f"Generated: {meta.get('generated_at', '?')} | "
            f"Model: {meta.get('model', '?')} | "
            f"Recordings: {meta.get('recordings_processed', '?')} | "
            f"Version: {meta.get('version', '?')}"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(150, 150, 150)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# Daily Report Compat Layer
# ============================================================
# Maps meeting minutes schema → daily report schema so the
# existing frontend can render meeting data with zero changes.
#
# Field mapping:
#   category: strategy|operations|finance|partnership|hr|legal|general → progress
#             product|technical → quality
#   action_items.owner → action_items.responsible
#   key_decisions: [{decision, rationale, decided_by}] → ["decision (by X)"]
#   follow_ups + next_steps → synthetic final topic
#   safety_flags, safety_observations → empty []
#   open_questions → appended to summary text
# ============================================================

# Meeting category → daily report category
_CATEGORY_MAP = {
    'strategy': 'progress',
    'operations': 'progress',
    'finance': 'progress',
    'partnership': 'progress',
    'hr': 'progress',
    'legal': 'progress',
    'general': 'progress',
    'product': 'quality',
    'technical': 'quality',
}


def convert_to_daily_report_format(minutes, meeting_config, transcripts):
    """
    Convert meeting minutes JSON → daily report JSON compatible with
    the existing FieldSight frontend (reports/{date}/{user}/daily_report.json).

    Returns: (compat_report_dict, user_name_for_path)
    """
    target_date = minutes.get('meeting_date', meeting_config.get('date', ''))
    meeting_title = minutes.get('meeting_title', 'Meeting')
    attendees = minutes.get('attendees', [])

    # Determine user path — use first device found (report goes under that user folder)
    device = None
    if transcripts:
        device = transcripts[0].get('device', None)
    user_for_path = meeting_config.get('user') or device or 'meeting'

    # Convert topics
    compat_topics = []
    for topic in minutes.get('topics', []):
        # Map category
        raw_cat = topic.get('category', 'general')
        compat_cat = _CATEGORY_MAP.get(raw_cat, 'progress')

        # Flatten key_decisions: [{decision, rationale, decided_by}] → ["text"]
        raw_decisions = topic.get('key_decisions', [])
        flat_decisions = []
        for d in raw_decisions:
            if isinstance(d, str):
                flat_decisions.append(d)
            elif isinstance(d, dict):
                text = d.get('decision', '')
                by = d.get('decided_by', '')
                if by:
                    text += f" (by {by})"
                flat_decisions.append(text)

        # Map action_items: owner → responsible
        compat_actions = []
        for ai in topic.get('action_items', []):
            compat_actions.append({
                'action': ai.get('action', ''),
                'responsible': ai.get('owner', ai.get('responsible', '?')),
                'deadline': ai.get('deadline', '?'),
                'priority': ai.get('priority', 'medium'),
            })

        # Append open questions to summary
        summary = topic.get('summary', '')
        open_qs = topic.get('open_questions', [])
        if open_qs:
            summary += ' Open questions: ' + '; '.join(open_qs)

        compat_topics.append({
            'topic_id': topic.get('topic_id', 0),
            'time_range': topic.get('time_range', ''),
            'topic_title': topic.get('topic_title', ''),
            'category': compat_cat,
            'summary': summary,
            'participants': topic.get('participants', []),
            'key_decisions': flat_decisions,
            'action_items': compat_actions,
            'safety_flags': [],
            'related_photos': [],
        })

    # Synthetic topic for follow_ups + next_steps + parking_lot
    follow_ups = minutes.get('follow_ups', [])
    next_steps = minutes.get('next_steps', [])
    parking_lot = minutes.get('parking_lot', [])

    if follow_ups or next_steps or parking_lot:
        synth_actions = []
        for fu in follow_ups:
            synth_actions.append({
                'action': fu.get('item', ''),
                'responsible': fu.get('owner', '?'),
                'deadline': fu.get('deadline', '?'),
                'priority': fu.get('priority', 'medium'),
            })

        synth_decisions = list(next_steps)
        synth_summary = ''
        if parking_lot:
            synth_summary = 'Parked for future discussion: ' + '; '.join(parking_lot)

        next_id = max((t.get('topic_id', 0) for t in compat_topics), default=-1) + 1
        compat_topics.append({
            'topic_id': next_id,
            'time_range': '',
            'topic_title': 'Follow-ups & Next Steps',
            'category': 'progress',
            'summary': synth_summary,
            'participants': attendees,
            'key_decisions': synth_decisions,
            'action_items': synth_actions,
            'safety_flags': [],
            'related_photos': [],
        })

    compat_report = {
        'report_date': target_date,
        'report_type': 'daily',
        'user_name': user_for_path,
        'device': device or 'meeting',
        'site': meeting_title,  # Use meeting title as "site" — shows in UI header
        'executive_summary': minutes.get('executive_summary', ''),
        'topics': compat_topics,
        'safety_observations': [],
        '_report_metadata': {
            **minutes.get('_report_metadata', {}),
            'source': 'meeting_minutes',
            'meeting_title': meeting_title,
            'meeting_type': minutes.get('meeting_type', 'general'),
        }
    }

    return compat_report, user_for_path


# ============================================================
# Generate Meeting Minutes
# ============================================================

def generate_meeting_minutes(meeting_config):
    """
    Main generation flow.

    Args:
        meeting_config: dict with keys:
            date, meeting_title, meeting_type, attendees,
            user (optional filter), transcript_prefix (optional custom path)
    """
    target_date = meeting_config['date']
    meeting_title = meeting_config.get('meeting_title', f'Meeting — {target_date}')
    meeting_type = meeting_config.get('meeting_type', 'general')
    attendees = meeting_config.get('attendees', [])

    logger.info(f"=== Generating meeting minutes: {meeting_title} ({target_date}) ===")

    # Collect transcripts
    transcripts = collect_transcripts(
        S3_BUCKET, target_date,
        user_filter=meeting_config.get('user'),
        custom_prefix=meeting_config.get('transcript_prefix'),
    )

    if not transcripts:
        logger.warning(f"No transcripts found for {target_date}")
        return {
            'status': 'no_data',
            'date': target_date,
            'meeting_title': meeting_title,
            'transcripts_found': 0,
        }

    total_words = sum(t.get('word_count', 0) for t in transcripts)

    # Auto-discover attendees from transcripts if not provided
    if not attendees:
        seen = set()
        for t in transcripts:
            name = t.get('speaker_name', t.get('device', ''))
            if name and name not in seen:
                attendees.append(name)
                seen.add(name)
        meeting_config['attendees'] = attendees

    # Build prompt
    prompt = build_meeting_prompt(transcripts, meeting_config)

    # Dynamic max_tokens: longer meetings need more output space
    # ~26K token input for 2hr meeting → expect 8-16K token output
    prompt_tokens_est = len(prompt) // 4
    max_tokens = min(max(8000, prompt_tokens_est // 2), 16000)
    logger.info(f"  Prompt ~{prompt_tokens_est} tokens → max_tokens={max_tokens}")

    # Call Claude
    raw_response, error = call_claude_structured(prompt, max_tokens=max_tokens)

    if error:
        logger.error(f"Claude error: {error}")
        save_debug_record(
            S3_BUCKET, target_date, meeting_title, prompt, None, None,
            parse_success=False,
            input_stats={
                'transcripts_count': len(transcripts),
                'total_words': total_words,
                'error': error,
            }
        )
        return {
            'status': 'error',
            'date': target_date,
            'meeting_title': meeting_title,
            'error': error,
        }

    # Parse JSON
    claude_output = extract_json_from_response(raw_response)
    parse_success = claude_output is not None

    if not claude_output:
        logger.error("Failed to parse Claude JSON, saving raw text as fallback")
        claude_output = {
            'executive_summary': raw_response[:500] if raw_response else 'Error generating minutes',
            'topics': [],
            'follow_ups': [],
            'next_steps': [],
            'parking_lot': [],
        }

    # Save debug
    save_debug_record(
        S3_BUCKET, target_date, meeting_title, prompt, raw_response, claude_output,
        parse_success=parse_success,
        input_stats={
            'transcripts_count': len(transcripts),
            'total_words': total_words,
            'attendees': attendees,
        }
    )

    # Build final minutes document
    now_iso = datetime.utcnow().isoformat() + 'Z'
    safe_title = re.sub(r'[^a-zA-Z0-9_-]', '_', meeting_title)[:50]

    minutes = {
        'meeting_date': target_date,
        'meeting_title': meeting_title,
        'meeting_type': meeting_type,
        'attendees': attendees,
        'executive_summary': claude_output.get('executive_summary', ''),
        'topics': claude_output.get('topics', []),
        'follow_ups': claude_output.get('follow_ups', []),
        'next_steps': claude_output.get('next_steps', []),
        'parking_lot': claude_output.get('parking_lot', []),
        '_report_metadata': {
            'version': 'v1.1',
            'generated_at': now_iso,
            'generated_by': meeting_config.get('triggered_by', 'system'),
            'recordings_processed': len(transcripts),
            'total_words': total_words,
            'model': CLAUDE_MODEL,
            'parse_success': parse_success,
        }
    }

    # Save JSON to S3 — archive copy in meeting_minutes/ prefix
    json_key = f"{MINUTES_PREFIX}{target_date}/{safe_title}.json"
    s3_client.put_object(
        Bucket=S3_BUCKET, Key=json_key,
        Body=json.dumps(minutes, ensure_ascii=False, indent=2, default=str),
        ContentType='application/json'
    )
    logger.info(f"Saved: {json_key}")

    # --- Determine user path for reports/ folder ---
    device = transcripts[0].get('device') if transcripts else None
    user_for_path = meeting_config.get('user') or device or 'meeting'

    # --- Save meeting_minutes.json to reports/ path (alongside daily_report) ---
    report_json_key = f"{REPORT_PREFIX}{target_date}/{user_for_path}/meeting_minutes.json"
    s3_client.put_object(
        Bucket=S3_BUCKET, Key=report_json_key,
        Body=json.dumps(minutes, ensure_ascii=False, indent=2, default=str),
        ContentType='application/json'
    )
    logger.info(f"Saved: {report_json_key}")

    # --- Save Word doc to both paths ---
    try:
        word_buffer = generate_word_document(
            minutes, f"Meeting Minutes — {meeting_title} — {target_date}"
        )
        if word_buffer:
            # Archive copy
            word_key = f"{MINUTES_PREFIX}{target_date}/{safe_title}.docx"
            s3_client.put_object(
                Bucket=S3_BUCKET, Key=word_key,
                Body=word_buffer.getvalue(),
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            # Reports/ path copy
            word_buffer.seek(0)
            report_word_key = f"{REPORT_PREFIX}{target_date}/{user_for_path}/meeting_minutes.docx"
            s3_client.put_object(
                Bucket=S3_BUCKET, Key=report_word_key,
                Body=word_buffer.getvalue(),
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            logger.info(f"Saved: {report_word_key}")
    except Exception as e:
        logger.error(f"Word generation failed: {e}")

    # --- Compat: write daily_report.json format for frontend backward compat ---
    compat_key = None
    try:
        compat_report, _ = convert_to_daily_report_format(
            minutes, meeting_config, transcripts
        )
        compat_key = f"{REPORT_PREFIX}{target_date}/{user_for_path}/daily_report.json"
        s3_client.put_object(
            Bucket=S3_BUCKET, Key=compat_key,
            Body=json.dumps(compat_report, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json'
        )
        logger.info(f"Saved compat report: {compat_key}")
    except Exception as e:
        logger.error(f"Compat report save failed: {e}")

    # --- Write meeting manifest (marks transcripts as consumed) ---
    manifest_key = None
    try:
        consumed_keys = [t.get('key', '') for t in transcripts if t.get('key')]
        manifest_key = write_meeting_manifest(
            s3_client, S3_BUCKET, REPORT_PREFIX, target_date,
            user_for_path, consumed_keys, meeting_title=meeting_title,
        )
        logger.info(f"Saved manifest: {manifest_key} ({len(consumed_keys)} keys)")
    except Exception as e:
        logger.error(f"Manifest write failed: {e}")

    return {
        'status': 'success',
        'date': target_date,
        'meeting_title': meeting_title,
        'transcripts_processed': len(transcripts),
        'total_words': total_words,
        'topics_extracted': len(minutes.get('topics', [])),
        'action_items': sum(
            len(t.get('action_items', []))
            for t in minutes.get('topics', [])
        ),
        's3_key': json_key,
        'report_key': report_json_key,
        'compat_report_key': compat_key,
        'manifest_key': manifest_key,
    }


# ============================================================
# MAIN HANDLER
# ============================================================

def lambda_handler(event, context):
    """
    Main Lambda handler.

    Event payload options:
      {"date": "2026-03-20"}
      {"date": "2026-03-20", "meeting_title": "BD Brainstorm", "meeting_type": "brainstorm"}
      {"date": "2026-03-20", "attendees": ["Ben", "Benny", "Sam"]}
      {"date": "2026-03-20", "user": "Benl1"}
      {"date": "2026-03-20", "transcript_prefix": "meetings/2026-03-20/session1/"}
      {"date": "today"}
      {"date": "yesterday"}
    """
    logger.info("=" * 60)
    logger.info("Meeting Minutes Generator v1.1 - Starting")
    logger.info(f"Event: {json.dumps(event, default=str)}")
    logger.info(f"Word generation: {'enabled' if DOCX_AVAILABLE else 'DISABLED'}")
    logger.info("=" * 60)

    if not S3_BUCKET:
        return {'statusCode': 400, 'body': 'Missing S3_BUCKET'}

    # Reset cache
    global _user_mapping_cache
    _user_mapping_cache = None

    # Build meeting config from event
    meeting_config = {
        'date': resolve_date(event.get('date', 'today')),
        'meeting_title': event.get('meeting_title', ''),
        'meeting_type': event.get('meeting_type', 'general'),
        'attendees': event.get('attendees', []),
        'user': event.get('user', None),
        'transcript_prefix': event.get('transcript_prefix', None),
        'triggered_by': event.get('triggered_by', 'system'),
    }

    # Default title if not provided
    if not meeting_config['meeting_title']:
        meeting_config['meeting_title'] = f"Meeting — {meeting_config['date']}"

    result = generate_meeting_minutes(meeting_config)

    logger.info("=" * 60)
    logger.info(f"Meeting Minutes Complete: {json.dumps(result, default=str)}")
    logger.info("=" * 60)

    return {
        'statusCode': 200,
        'body': json.dumps(result, default=str)
    }
