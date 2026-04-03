"""
Lambda 5: Daily/Weekly/Monthly Report Generator v3.5

Changes from v3.3:
- REFACTOR: Transcript parsing delegated to shared transcript_utils.py
- ADD: Meeting manifest reading — skips transcripts consumed by meeting minutes
- ADD: read_meeting_manifest() excludes meeting transcripts from daily site reports
- ADD: Segment end time in build_daily_prompt for accurate time_range
- CHANGE: Default CLAUDE_MODEL to claude-sonnet-4-6
- CHANGE: Transcript text truncation 15K → 60K chars in daily prompt

Changes from v3.2:
- CHANGE: extract_timestamp_from_filename() adds VAD offset for absolute time
- ADD: extract_vad_info_from_filename() returns source format and video availability
- ADD: transcript entry includes 'vad' field with offset/source metadata

Changes from v3.1:
- ADD: Per-user weekly/monthly reports (each user gets their own weekly summary)
- ADD: Per-site weekly/monthly reports (aggregate by site from user_mapping v2)
- ADD: get_user_site_mapping() to resolve user → site from v2 mapping format
- CHANGE: build_weekly_prompt now accepts scope_label/scope_type for targeted prompts
- CHANGE: generate_periodic_report produces 3 levels: per-user, per-site, combined
- CHANGE: Prefers per-user daily reports over summary when building weekly reports

Changes from v2:
- FIX: user_mapping v2 format (nested objects with name/role/site)
- ADD: Debug file saving (prompt + raw response + parsed JSON) for prompt tuning
- ADD: Stale report detection + automatic backfill for past 7 days
- ADD: python-docx Lambda Layer check with clear warning
- FIX: Timeout on Claude API increased to 180s

This Lambda function:
1. Collects all transcripts from target period (per user)
2. Collects all photos from target period (per user)
3. Correlates photos with transcript timestamps
4. Calls Claude API to generate STRUCTURED topic-based JSON (not free-text)
5. Saves debug record (prompt + response) for prompt tuning
6. Writes report JSON + Word doc + metadata to S3 and DynamoDB
7. Auto-detects stale reports from past 7 days and regenerates
8. Supports daily / weekly / monthly report modes

Trigger:
  - EventBridge daily   at 05:00 NZDT  → {"report_type": "daily"}
  - EventBridge weekly  Fri 18:00 NZDT → {"report_type": "weekly"}
  - EventBridge monthly 1st 06:00 NZDT → {"report_type": "monthly"}
  - API Gateway (on-demand regenerate)  → {"report_type": "daily", "date": "2026-02-19", "hidden_topic_ids": [3]}

Environment Variables:
    S3_BUCKET           - S3 bucket name
    ANTHROPIC_API_KEY   - Anthropic API key (sk-ant-xxx)
    CLAUDE_MODEL        - Claude model ID (default: claude-sonnet-4-6)
    REPORT_PREFIX       - Report output prefix (default: reports/)
    ITEMS_TABLE         - DynamoDB table for items (default: fieldsight-items)
    REPORTS_TABLE       - DynamoDB table for reports (default: fieldsight-reports)
    AUDIT_TABLE         - DynamoDB table for audit log (default: fieldsight-audit)
    PROMPT_TEMPLATES_KEY - S3 key for prompt templates (default: config/prompt_templates.json)
    BACKFILL_DAYS       - How many days back to check for stale reports (default: 7)
"""

import os
import json
import logging
import re
import boto3
from boto3.dynamodb.conditions import Key
import urllib3
from datetime import datetime, timedelta
from io import BytesIO
from transcript_utils import (
    normalize_transcript, format_turns_for_prompt, get_time_bounds,
    extract_device_from_filename as tu_extract_device,
    extract_vad_metadata_from_filename as tu_extract_vad_info,
    read_meeting_manifest,
)

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
s3_client = boto3.client('s3')
dynamodb = boto3.resource('dynamodb')

# Configuration
S3_BUCKET = os.environ.get('S3_BUCKET', '')
REPORT_PREFIX = os.environ.get('REPORT_PREFIX', 'reports/')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
CLAUDE_MODEL = os.environ.get('CLAUDE_MODEL', 'claude-sonnet-4-6')
ITEMS_TABLE = os.environ.get('ITEMS_TABLE', 'fieldsight-items')
REPORTS_TABLE = os.environ.get('REPORTS_TABLE', 'fieldsight-reports')
AUDIT_TABLE = os.environ.get('AUDIT_TABLE', 'fieldsight-audit')
BACKFILL_DAYS = int(os.environ.get('BACKFILL_DAYS', '7'))
ENABLE_DYNAMODB = os.environ.get('ENABLE_DYNAMODB', 'false').lower() == 'true'
PROMPT_TEMPLATES_KEY = os.environ.get('PROMPT_TEMPLATES_KEY', 'config/prompt_templates.json')
CORRECTIONS_TABLE = os.environ.get('CORRECTIONS_TABLE', 'fieldsight-corrections')

# Prompt template cache (loaded once per invocation)
_prompt_templates_cache = None


def load_prompt_templates(bucket):
    global _prompt_templates_cache
    if _prompt_templates_cache is not None:
        return _prompt_templates_cache
    try:
        resp = s3_client.get_object(Bucket=bucket, Key=PROMPT_TEMPLATES_KEY)
        templates = json.loads(resp['Body'].read().decode('utf-8'))
        _prompt_templates_cache = templates
        logger.info(f"Loaded prompt templates v{templates.get('_version', '?')} "
                    f"from s3://{bucket}/{PROMPT_TEMPLATES_KEY}")
        return templates
    except Exception as e:
        logger.warning(f"Prompt templates not found ({e}) — using inline defaults")
        _prompt_templates_cache = {}
        return {}


def get_template(template_name, field='prompt'):
    templates = _prompt_templates_cache or {}
    section = templates.get(template_name, {})
    return section.get(field)


# ============================================================
# Corrections helpers (QA/QC Layer 2)
# ============================================================

def fetch_corrections_for_range(start_date, end_date):
    """Fetch all active corrections for a date range from DynamoDB.
    Returns dict: {date: {(topic_id, field): correction_item}}
    """
    corrections = {}
    try:
        table = dynamodb.Table(CORRECTIONS_TABLE)
        for date_str in dates_in_range(start_date, end_date):
            resp = table.query(
                KeyConditionExpression=Key('PK').eq(f'DATE#{date_str}') & Key('SK').begins_with('CORRECTION#')
            )
            items = resp.get('Items', [])
            while resp.get('LastEvaluatedKey'):
                resp = table.query(
                    KeyConditionExpression=Key('PK').eq(f'DATE#{date_str}') & Key('SK').begins_with('CORRECTION#'),
                    ExclusiveStartKey=resp['LastEvaluatedKey']
                )
                items.extend(resp.get('Items', []))

            for item in items:
                if item.get('status') != 'active':
                    continue
                tid = int(item.get('topic_id', -1))
                field = item.get('field', '')
                key = (tid, field)
                date_corrections = corrections.setdefault(date_str, {})
                existing = date_corrections.get(key)
                if not existing or item.get('submitted_at', '') > existing.get('submitted_at', ''):
                    date_corrections[key] = item
    except Exception as e:
        logger.warning(f"Failed to fetch corrections: {e}")
    return corrections


def apply_corrections_to_reports(daily_reports, corrections_map):
    """Apply corrections to daily report data in-place. Returns count applied."""
    total_applied = 0
    for report in daily_reports:
        date = report.get('report_date', '')
        date_corrections = corrections_map.get(date, {})
        if not date_corrections:
            continue
        for topic in report.get('topics', []):
            topic_id = topic.get('topic_id')
            for (corr_tid, corr_field), corr_item in date_corrections.items():
                if corr_tid == topic_id and topic.get(corr_field) is not None:
                    topic[corr_field] = corr_item['corrected']
                    total_applied += 1
    return total_applied


# Check python-docx availability once at import time
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
    logger.info("python-docx available — Word generation enabled")
except ImportError:
    logger.warning(
        "python-docx NOT installed — Word generation disabled. "
        "To enable: add a Lambda Layer with python-docx. "
        "See: https://docs.aws.amazon.com/lambda/latest/dg/chapter-layers.html"
    )


# ============================================================
# JSON Schema (what Claude MUST output)
# ============================================================
DAILY_REPORT_SCHEMA = """{
  "executive_summary": "2-3 sentence overview of key activities and outcomes",
  "quality_and_compliance": [
    {
      "item": "Testing or compliance matter",
      "status": "pending | in_progress | completed | concern",
      "details": "Specific requirements mentioned",
      "follow_up_needed": true
    }
  ],
  "safety_observations": [
    {
      "observation": "What was observed",
      "risk_level": "high | medium | low",
      "location": "Where on site",
      "who_raised": "Speaker name",
      "recommended_action": "What should be done"
    }
  ],
  "critical_dates_and_deadlines": [
    {
      "date_mentioned": "Date or timeframe from conversation",
      "context": "What happens on that date",
      "who_mentioned": "Speaker name",
      "urgency": "high | medium | low",
      "type": "deadline | inspection | delivery | weather | meeting | other"
    }
  ],
  "topics": [
    {
      "topic_id": 0,
      "time_range": "HH:MM – HH:MM",
      "topic_title": "Short descriptive title",
      "category": "safety | progress | quality",
      "participants": ["Name1", "Name2"],
      "summary": "2-4 sentence summary of what was discussed and decided",
      "key_decisions": ["Decision attributed to person"],
      "action_items": [
        {
          "action": "What needs to be done",
          "responsible": "Person name",
          "deadline": "When (e.g. 'Tomorrow 08:00', 'EOD', '15:00'), or null if not mentioned",
          "priority": "high | medium | low"
        }
      ],
      "safety_flags": [
        {
          "observation": "What was observed",
          "risk_level": "high | medium | low",
          "recommended_action": "What should be done"
        }
      ],
      "related_photos": ["filename.jpg"]
    }
  ]
}"""

WEEKLY_REPORT_SCHEMA = """{
  "executive_summary": "3-5 sentence overview of the week",
  "safety_trends": [
    {
      "trend": "Description of safety trend",
      "risk_level": "high | medium | low",
      "frequency": "How often observed this week",
      "recommendation": "What to do about it"
    }
  ],
  "progress_highlights": [
    {
      "item": "What was accomplished",
      "date": "YYYY-MM-DD",
      "status": "completed | in_progress | delayed"
    }
  ],
  "outstanding_actions": [
    {
      "action": "What needs to be done",
      "responsible": "Person name",
      "original_date": "YYYY-MM-DD",
      "priority": "high | medium | low",
      "status": "open | overdue | completed"
    }
  ],
  "quality_summary": "Overview of quality observations this week",
  "next_week_priorities": ["Priority 1", "Priority 2"]
}"""


# ============================================================
# Date Helpers
# ============================================================

def get_nzdt_now():
    return datetime.utcnow() + timedelta(hours=13)

def get_yesterday_date():
    return (get_nzdt_now() - timedelta(days=1)).strftime('%Y-%m-%d')

def get_week_range():
    nzdt_now = get_nzdt_now()
    days_since_sunday = (nzdt_now.weekday() - 6) % 7
    if days_since_sunday == 0:
        days_since_sunday = 7
    last_sunday = nzdt_now - timedelta(days=days_since_sunday)
    last_monday = last_sunday - timedelta(days=6)
    return last_monday.strftime('%Y-%m-%d'), last_sunday.strftime('%Y-%m-%d')

def get_month_range():
    nzdt_now = get_nzdt_now()
    first_of_this_month = nzdt_now.replace(day=1)
    last_of_prev_month = first_of_this_month - timedelta(days=1)
    first_of_prev_month = last_of_prev_month.replace(day=1)
    return first_of_prev_month.strftime('%Y-%m-%d'), last_of_prev_month.strftime('%Y-%m-%d')

def dates_in_range(start_str, end_str):
    start = datetime.strptime(start_str, '%Y-%m-%d')
    end = datetime.strptime(end_str, '%Y-%m-%d')
    dates = []
    cur = start
    while cur <= end:
        dates.append(cur.strftime('%Y-%m-%d'))
        cur += timedelta(days=1)
    return dates


# ============================================================
# S3 Helpers
# ============================================================

def list_s3_objects(bucket, prefix):
    objects = []
    paginator = s3_client.get_paginator('list_objects_v2')
    for page in paginator.paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get('Contents', []):
            objects.append({
                'key': obj['Key'],
                'last_modified': obj['LastModified'],
                'size': obj['Size']
            })
    return objects

def download_json_from_s3(bucket, key):
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read().decode('utf-8')
        return json.loads(content)
    except Exception as e:
        logger.error(f"Failed to download {key}: {str(e)}")
        return None


# ============================================================
# User Mapping
# ============================================================

_user_mapping_cache = None

def load_user_mapping(bucket):
    global _user_mapping_cache
    if _user_mapping_cache is not None:
        return _user_mapping_cache
    try:
        data = download_json_from_s3(bucket, 'config/user_mapping.json')
        if data:
            raw_mapping = data.get('mapping', {})
            normalized = {}
            for device, value in raw_mapping.items():
                if isinstance(value, str):
                    normalized[device] = value
                elif isinstance(value, dict):
                    normalized[device] = value.get('name', device)
                else:
                    normalized[device] = str(value)
            _user_mapping_cache = normalized
            logger.info(f"Loaded user mapping: {len(normalized)} entries "
                        f"(format: {'v2' if any(isinstance(v, dict) for v in raw_mapping.values()) else 'v1'})")
            return _user_mapping_cache
    except Exception as e:
        logger.warning(f"Failed to load user mapping: {e}")
    _user_mapping_cache = {}
    return _user_mapping_cache

def load_user_mapping_full(bucket):
    try:
        data = download_json_from_s3(bucket, 'config/user_mapping.json')
        if data:
            return {'mapping': data.get('mapping', {}), 'sites': data.get('sites', {})}
    except:
        pass
    return {'mapping': {}, 'sites': {}}

def get_user_site_mapping(bucket):
    full = load_user_mapping_full(bucket)
    raw_mapping = full.get('mapping', {})
    sites_info = full.get('sites', {})
    user_primary_site = {}
    user_all_sites = {}
    user_roles = {}
    for device, value in raw_mapping.items():
        if isinstance(value, dict):
            name = value.get('name', device)
            primary = value.get('primary_site', '')
            all_sites = value.get('sites', [])
            role = value.get('role', '')
            name_variants = {name, name.replace(' ', '_'), name.replace('_', ' ')}
            for n in name_variants:
                if primary:
                    user_primary_site[n] = primary
                if all_sites:
                    user_all_sites[n] = all_sites
                elif primary:
                    user_all_sites[n] = [primary]
                if role:
                    user_roles[n] = role
    return user_primary_site, user_all_sites, user_roles, sites_info


# ============================================================
# Transcript Parsing — delegates to transcript_utils.py
# ============================================================

def extract_timestamp_from_filename(filename):
    """Extract timestamp with VAD offset — delegates to transcript_utils."""
    from transcript_utils import compute_segment_base_time
    return compute_segment_base_time(filename)

def extract_device_from_filename(filename):
    return tu_extract_device(filename)

def extract_vad_info_from_filename(filename):
    return tu_extract_vad_info(filename)

def parse_transcript(transcript_data):
    """Parse AWS Transcribe JSON — delegates to transcript_utils."""
    from transcript_utils import parse_transcribe_json
    parsed = parse_transcribe_json(transcript_data)
    if not parsed:
        return None
    return {
        'full_text': parsed['full_text'],
        'words': parsed['words'],
        'word_count': parsed['word_count'],
        'duration_seconds': parsed['duration_seconds'],
    }

def correlate_photos_with_transcripts(transcripts, photos):
    correlated = []
    for transcript in transcripts:
        t_time = transcript.get('timestamp')
        related = []
        if t_time:
            for photo in photos:
                p_time = photo.get('timestamp')
                if p_time and abs((t_time - p_time).total_seconds()) <= 300:
                    related.append({
                        'key': photo['key'],
                        'filename': photo.get('filename', ''),
                        'time_diff_seconds': abs((t_time - p_time).total_seconds()),
                        'time_str': photo.get('time_str', '')
                    })
        related.sort(key=lambda x: x.get('time_diff_seconds', 999))
        correlated.append({**transcript, 'related_photos': related[:5]})
    return correlated


# ============================================================
# Claude API
# ============================================================

def call_claude_structured(prompt, max_tokens=4096):
    if not ANTHROPIC_API_KEY:
        logger.error("ANTHROPIC_API_KEY not set")
        return None, "ANTHROPIC_API_KEY not configured"
    http = urllib3.PoolManager()
    body = json.dumps({
        "model": CLAUDE_MODEL,
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}]
    })
    try:
        resp = http.request(
            'POST', 'https://api.anthropic.com/v1/messages',
            body=body,
            headers={
                'Content-Type': 'application/json',
                'x-api-key': ANTHROPIC_API_KEY,
                'anthropic-version': '2023-06-01',
            },
            timeout=180.0,
        )
        data = json.loads(resp.data.decode('utf-8'))
        if resp.status == 200:
            text_blocks = [b['text'] for b in data.get('content', []) if b.get('type') == 'text']
            return '\n'.join(text_blocks), None
        else:
            err = data.get('error', {}).get('message', f'HTTP {resp.status}')
            logger.error(f"Claude API error: {err}")
            return None, err
    except Exception as e:
        logger.error(f"Claude API call failed: {str(e)}")
        return None, str(e)

def extract_json_from_response(raw_text):
    match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw_text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    try:
        return json.loads(raw_text.strip())
    except json.JSONDecodeError:
        pass
    first_brace = raw_text.find('{')
    last_brace = raw_text.rfind('}')
    if first_brace != -1 and last_brace != -1:
        try:
            return json.loads(raw_text[first_brace:last_brace + 1])
        except json.JSONDecodeError:
            pass
    logger.error(f"Failed to extract JSON from Claude response: {raw_text[:500]}")
    return None


# ============================================================
# Debug Record
# ============================================================

def save_debug_record(bucket, target_date, user_name, prompt, raw_response,
                      parsed_json, parse_success, input_stats):
    try:
        if user_name == '_summary':
            debug_key = f"{REPORT_PREFIX}{target_date}/summary_report_debug.json"
        else:
            debug_key = f"{REPORT_PREFIX}{target_date}/{user_name}/daily_report_debug.json"
        debug_record = {
            '_description': 'Debug record for prompt tuning.',
            'timestamp': datetime.utcnow().isoformat() + 'Z',
            'model': CLAUDE_MODEL,
            'user_name': user_name,
            'target_date': target_date,
            'parse_success': parse_success,
            'input_stats': input_stats,
            'prompt': prompt,
            'prompt_length': len(prompt),
            'raw_response': raw_response,
            'raw_response_length': len(raw_response) if raw_response else 0,
            'parsed_json': parsed_json,
        }
        s3_client.put_object(
            Bucket=bucket, Key=debug_key,
            Body=json.dumps(debug_record, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json'
        )
        logger.info(f"  Saved debug: {debug_key}")
    except Exception as e:
        logger.error(f"  Debug save failed for {user_name}: {e}")


# ============================================================
# Prompt Builders
# ============================================================

def build_daily_prompt(transcripts_with_photos, user_name, site_name, target_date,
                       role=None, total_duration=0.0, num_photos=0, name_mapping=None):
    """Build the Claude prompt for daily structured report."""
    transcript_lines = []
    for t in transcripts_with_photos:
        time_str = t.get('time_str', '??:??')
        device = t.get('device', 'Unknown')
        text = t.get('text', '')
        photos = t.get('related_photos', [])
        photo_names = [os.path.basename(p.get('filename', p.get('key', ''))) for p in photos]

        # Compute segment end time from start + duration for accurate time_range
        seg_duration = t.get('duration_seconds', 0)
        ts = t.get('timestamp')
        if ts and seg_duration > 0:
            end_time = ts + timedelta(seconds=seg_duration)
            time_label = f"{ts.strftime('%H:%M:%S')} \u2013 {end_time.strftime('%H:%M:%S')}"
        else:
            time_label = time_str

        line = f"[{time_label}] {device}: {text}"
        if photo_names:
            line += f"\n  Photos taken near this time: {', '.join(photo_names)}"
        transcript_lines.append(line)

    transcripts_text = "\n\n".join(transcript_lines)[:60000]

    dur_min = int(total_duration // 60)
    dur_sec = int(total_duration % 60)
    dur_str = f"{dur_min}m {dur_sec}s" if dur_min > 0 else f"{dur_sec}s"

    meta_lines = [
        f"- Date: {target_date}",
        f"- Site: {site_name}",
        f"- Worker: {user_name}",
    ]
    if role:
        meta_lines.append(f"- Role: {role}")
    meta_lines.append(f"- Recordings: {len(transcripts_with_photos)} ({dur_str} total audio)")
    if num_photos:
        meta_lines.append(f"- Photos: {num_photos}")
    metadata_block = "\n".join(meta_lines)

    name_ref_block = ""
    if name_mapping:
        devices_used = set(t.get('device', '') for t in transcripts_with_photos)
        ref_lines = []
        for dev in sorted(devices_used):
            display = name_mapping.get(dev, dev)
            if display != dev:
                ref_lines.append(f"- {dev} = {display}")
        if ref_lines:
            name_ref_block = "\n## Name Reference (device account \u2192 person name)\n" + "\n".join(ref_lines)
            name_ref_block += "\nIMPORTANT: Always use the person's real name in participants, who_raised, who_mentioned, and responsible fields. Never use device account names (Benl1, Benl2, etc.)."

    template = get_template('daily_report', 'prompt')
    system_ctx = get_template('daily_report', 'system_context') or \
        "You are a construction site documentation assistant for a New Zealand construction company."

    if template:
        body = template.format(
            metadata_block=metadata_block,
            name_reference=name_ref_block,
            transcripts_text=transcripts_text,
            schema=DAILY_REPORT_SCHEMA,
        )
        return f"{system_ctx}\n\n{body}"

    return f"""{system_ctx}

## Recording Session Info
{metadata_block}
{name_ref_block}

Analyze the following radio transcript recordings and produce a STRUCTURED daily report.

## Raw Transcript Recordings (chronological)
{transcripts_text}

## Instructions
1. Write an executive_summary covering the whole day in 2-3 sentences.
2. Extract quality_and_compliance items: testing, inspections, code compliance, certifications mentioned.
3. Extract safety_observations: any hazards, near-misses, PPE issues, or safety concerns raised. Include who raised it and where.
4. Extract critical_dates_and_deadlines: any dates, deadlines, inspections, deliveries, or time-sensitive items mentioned.
5. Group the recordings into logical TOPICS (e.g. "Morning Safety Briefing", "Block C Concrete Pour", "Scaffold Inspection").
6. For each topic, classify as safety/progress/quality, list participants by name, and extract decisions and action items.
7. Photo filenames are provided where available \u2014 include them in the relevant topic's related_photos.

## Output Format
Return ONLY valid JSON matching this exact schema (no markdown, no explanation, just the JSON object):

{DAILY_REPORT_SCHEMA}

Rules:
- topic_id must be sequential starting from 0
- time_range format: "HH:MM \u2013 HH:MM" (use en-dash) \u2014 derive from the [HH:MM:SS \u2013 HH:MM:SS] timestamps shown in each transcript segment
- category MUST be one of: safety, progress, quality
- priority MUST be one of: high, medium, low
- risk_level MUST be one of: high, medium, low
- urgency MUST be one of: high, medium, low
- status MUST be one of: pending, in_progress, completed, concern
- type MUST be one of: deadline, inspection, delivery, weather, meeting, other
- participants should list names of people involved in each topic
- action_items, safety_flags, and all arrays can be empty []
- related_photos should contain filenames from the photo lines above
- Do NOT include any text outside the JSON object"""


def build_weekly_prompt(daily_reports, site_name, start_date, end_date,
                       scope_label=None, scope_type='site', user_role=None,
                       corrections_applied=0):
    summaries = []
    for report in daily_reports:
        date = report.get('report_date', '?')
        user = report.get('user_name', report.get('user', '?'))
        exec_sum = report.get('executive_summary', report.get('summary', 'No summary'))
        topics_text = ""
        for t in report.get('topics', []):
            cat = t.get('category', '?')
            title = t.get('topic_title', '?')
            flags = t.get('safety_flags', [])
            actions = t.get('action_items', [])
            topics_text += f"\n    - [{cat}] {title}"
            for f in flags:
                topics_text += f"\n      {f.get('observation', '')} (risk: {f.get('risk_level', '?')})"
            for a in actions:
                topics_text += f"\n      \u2610 {a.get('action', '')} \u2192 {a.get('responsible', '?')} by {a.get('deadline', '?')}"
        summaries.append(f"### {date} \u2014 {user}\n{exec_sum}{topics_text}")

    all_summaries = "\n\n".join(summaries)

    scope_intros = (_prompt_templates_cache or {}).get('scope_intros', {})

    if scope_type == 'user' and scope_label:
        role_context = f" (role: {user_role})" if user_role else ""
        intro_tpl = scope_intros.get('user') if scope_intros else None
        scope_intro = (
            intro_tpl.format(scope_label=scope_label, role_context=role_context,
                             start_date=start_date, end_date=end_date, site_name=site_name)
            if intro_tpl else
            f"Summarize the following daily reports for **{scope_label}**{role_context} "
            f"from {start_date} to {end_date} at site \"{site_name}\" "
            f"into a personal WEEKLY overview report."
        )
        extra_instructions = scope_intros.get('user_extra',
            "\n6. Focus on this individual's activities, responsibilities, and action items."
            "\n7. Highlight their personal safety observations and quality contributions."
        )
        if user_role in ('site_manager', 'pm'):
            extra_instructions += scope_intros.get('user_manager_extra',
                "\n8. As a manager/PM, emphasize supervisory decisions, team coordination, "
                "and items requiring their follow-up or escalation."
            )
    elif scope_type == 'site' and scope_label:
        intro_tpl = scope_intros.get('site') if scope_intros else None
        scope_intro = (
            intro_tpl.format(scope_label=scope_label, start_date=start_date, end_date=end_date)
            if intro_tpl else
            f"Summarize the following daily reports for site \"{scope_label}\" "
            f"from {start_date} to {end_date} into a site-level WEEKLY overview report."
        )
        extra_instructions = scope_intros.get('site_extra',
            "\n6. Aggregate across all workers on this site."
            "\n7. Identify cross-team coordination issues or dependencies."
        )
    else:
        intro_tpl = scope_intros.get('all') if scope_intros else None
        scope_intro = (
            intro_tpl.format(start_date=start_date, end_date=end_date, site_name=site_name)
            if intro_tpl else
            f"Summarize the following daily reports from {start_date} to {end_date} "
            f"at site \"{site_name}\" into a WEEKLY overview report."
        )
        extra_instructions = ""

    corrections_note = ""
    if corrections_applied > 0:
        corrections_note = (
            f"\n\n**Note:** {corrections_applied} user correction(s) have been applied to the "
            f"daily reports above. The corrected values are already reflected in the summaries. "
            f"Treat corrected content as authoritative."
        )

    template = get_template('weekly_report', 'prompt')
    system_ctx = get_template('weekly_report', 'system_context') or \
        "You are a construction site documentation assistant for a New Zealand construction company."

    if template:
        body = template.format(
            scope_intro=scope_intro,
            all_summaries=all_summaries[:15000],
            corrections_note=corrections_note,
            extra_instructions=extra_instructions,
            schema=WEEKLY_REPORT_SCHEMA,
        )
        return f"{system_ctx}\n\n{body}"

    return f"""{system_ctx}

{scope_intro}

## Daily Report Summaries
{all_summaries[:15000]}{corrections_note}

## Instructions
1. Identify safety TRENDS across the week (recurring issues, improvements, new risks).
2. List progress highlights with completion status.
3. Compile outstanding action items \u2014 mark overdue items.
4. Summarize quality observations.
5. Suggest priorities for next week.{extra_instructions}

## Output Format
Return ONLY valid JSON matching this exact schema:

{WEEKLY_REPORT_SCHEMA}

Rules:
- status MUST be one of the allowed values shown
- risk_level and priority MUST be one of: high, medium, low
- Do NOT include any text outside the JSON object"""


def build_monthly_prompt(daily_reports, weekly_reports, site_name, start_date, end_date,
                        corrections_applied=0):
    if weekly_reports:
        source_text = "\n\n".join([
            f"### Week of {r.get('period', {}).get('start', '?')} to {r.get('period', {}).get('end', '?')}\n"
            f"{r.get('executive_summary', 'No summary')}"
            for r in weekly_reports
        ])
        source_label = "Weekly Report Summaries"
    else:
        source_text = "\n\n".join([
            f"### {r.get('report_date', '?')}\n{r.get('executive_summary', r.get('summary', 'No summary'))}"
            for r in daily_reports
        ])
        source_label = "Daily Report Summaries"

    corrections_note = ""
    if corrections_applied > 0:
        corrections_note = (
            f"\n\n**Note:** {corrections_applied} user correction(s) have been applied to the "
            f"source reports. Treat corrected content as authoritative."
        )

    template = get_template('monthly_report', 'prompt')
    system_ctx = get_template('monthly_report', 'system_context') or \
        "You are a construction site documentation assistant for a New Zealand construction company."

    if template:
        body = template.format(
            site_name=site_name, start_date=start_date, end_date=end_date,
            source_label=source_label, source_text=source_text[:15000],
            corrections_note=corrections_note,
            schema=WEEKLY_REPORT_SCHEMA,
        )
        return f"{system_ctx}\n\n{body}"

    return f"""{system_ctx}

Summarize the following reports from {start_date} to {end_date} at site "{site_name}" into a MONTHLY overview.

## {source_label}
{source_text[:15000]}{corrections_note}

## Output Format
Return ONLY valid JSON matching this schema:

{WEEKLY_REPORT_SCHEMA}

Additional: Include month-level trends and patterns. "next_week_priorities" should be "next_month_priorities".

Rules:
- Do NOT include any text outside the JSON object"""


# ============================================================
# DynamoDB Helpers
# ============================================================

def write_items_to_dynamodb(site_id, target_date, topics, user_name, device):
    if not ENABLE_DYNAMODB:
        return
    try:
        table = dynamodb.Table(ITEMS_TABLE)
        pk = f"SITE#{site_id}#DATE#{target_date}"
        for topic in topics:
            tid = topic.get('topic_id', 0)
            time_range = topic.get('time_range', '00:00')
            start_time = time_range.split('\u2013')[0].strip().replace(':', '')
            sk = f"ITEM#{start_time}#{tid}"
            table.put_item(Item={
                'PK': pk, 'SK': sk, 'type': 'recording',
                'topic_id': tid,
                'topic_title': topic.get('topic_title', ''),
                'category': topic.get('category', 'progress'),
                'time_range': time_range,
                'participants': topic.get('participants', []),
                'summary': topic.get('summary', ''),
                'key_decisions': topic.get('key_decisions', []),
                'action_items': topic.get('action_items', []),
                'safety_flags': topic.get('safety_flags', []),
                'related_photos': topic.get('related_photos', []),
                'photo_count': len(topic.get('related_photos', [])),
                'user_name': user_name, 'device': device,
                'hidden': False, 'hidden_by': None, 'hidden_at': None,
            })
        logger.info(f"  DynamoDB: wrote {len(topics)} items to {ITEMS_TABLE}")
    except Exception as e:
        logger.error(f"  DynamoDB items write failed: {e}")

def write_report_to_dynamodb(site_id, target_date, report_type, s3_key, user_name,
                              visible_count, hidden_count):
    if not ENABLE_DYNAMODB:
        return
    try:
        table = dynamodb.Table(REPORTS_TABLE)
        now_iso = datetime.utcnow().isoformat() + 'Z'
        table.put_item(Item={
            'PK': f"SITE#{site_id}#DATE#{target_date}",
            'SK': f"REPORT#{report_type}#{now_iso}",
            'report_type': report_type, 's3_key': s3_key,
            'generated_by': 'system', 'generated_at': now_iso,
            'user_name': user_name,
            'visible_items': visible_count, 'hidden_items': hidden_count,
        })
        logger.info(f"  DynamoDB: wrote report metadata to {REPORTS_TABLE}")
    except Exception as e:
        logger.error(f"  DynamoDB report write failed: {e}")

def write_audit_entry(site_id, target_date, action, detail, user='System'):
    if not ENABLE_DYNAMODB:
        return
    try:
        table = dynamodb.Table(AUDIT_TABLE)
        now_iso = datetime.utcnow().isoformat() + 'Z'
        table.put_item(Item={
            'PK': f"SITE#{site_id}#DATE#{target_date}",
            'SK': f"AUDIT#{now_iso}",
            'action': action, 'user': user, 'detail': detail, 'timestamp': now_iso,
        })
    except Exception as e:
        logger.error(f"  DynamoDB audit write failed: {e}")


# ============================================================
# Word Document Generation
# ============================================================

def generate_word_document(report_data, title):
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Recording Session Info
    session = report_data.get('recording_session', {})
    if session:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light List'
        meta_rows = [('Date', session.get('date', '')), ('Site', session.get('site', ''))]
        if session.get('worker'):
            meta_rows.append(('Worker', session.get('worker', '')))
            role = session.get('role', '')
            if role:
                meta_rows.append(('Role', role.replace('_', ' ').title()))
        if session.get('workers'):
            meta_rows.append(('Workers', ', '.join(session.get('workers', []))))
        meta_rows.extend([
            ('Recordings', str(session.get('recordings', 0))),
            ('Total Audio', session.get('total_duration_display', '0m 0s')),
            ('Words Transcribed', f"{session.get('total_words', 0):,}"),
        ])
        if session.get('photos', 0) > 0:
            meta_rows.append(('Photos', str(session.get('photos', 0))))
        for label, value in meta_rows:
            if value:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(value)
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        doc.add_paragraph('')

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(report_data.get('executive_summary', 'No summary available'))

    quality_items = report_data.get('quality_and_compliance', [])
    if quality_items:
        doc.add_heading('Quality & Compliance', level=1)
        for qi in quality_items:
            status = qi.get('status', 'pending').upper()
            follow_up = 'Follow-up needed' if qi.get('follow_up_needed') else ''
            p = doc.add_paragraph()
            run = p.add_run(f"[{status}] ")
            run.bold = True
            if status == 'CONCERN':
                run.font.color.rgb = RGBColor(192, 57, 43)
            p.add_run(f"{qi.get('item', '')} \u2014 {qi.get('details', '')}")
            if follow_up:
                run2 = p.add_run(f"  {follow_up}")
                run2.font.color.rgb = RGBColor(211, 84, 0)

    safety_obs = report_data.get('safety_observations', [])
    if safety_obs:
        doc.add_heading('Safety Observations', level=1)
        for obs in safety_obs:
            risk = obs.get('risk_level', 'medium').upper()
            p = doc.add_paragraph()
            run = p.add_run(f"[{risk}] ")
            run.bold = True
            if risk == 'HIGH':
                run.font.color.rgb = RGBColor(192, 57, 43)
            who = obs.get('who_raised', '')
            location = obs.get('location', 'Unknown location')
            p.add_run(f"{obs.get('observation', '')} \u2014 {location}")
            if who:
                p.add_run(f" (raised by {who})")
            rec = obs.get('recommended_action', '')
            if rec:
                p.add_run(f"\n  \u2192 Recommended: {rec}")

    critical_dates = report_data.get('critical_dates_and_deadlines', [])
    if critical_dates:
        doc.add_heading('Critical Dates & Deadlines', level=1)
        for cd in critical_dates:
            urgency = cd.get('urgency', 'medium').upper()
            dtype = cd.get('type', 'other')
            p = doc.add_paragraph()
            run = p.add_run(f"[{urgency}] ")
            run.bold = True
            if urgency == 'HIGH':
                run.font.color.rgb = RGBColor(192, 57, 43)
            p.add_run(f"{cd.get('date_mentioned', '?')} \u2014 {cd.get('context', '')} ({dtype})")
            who = cd.get('who_mentioned', '')
            if who:
                p.add_run(f" \u2014 mentioned by {who}")

    topics = report_data.get('topics', [])
    if topics:
        doc.add_heading('Detailed Timeline', level=1)
        for topic in topics:
            cat = topic.get('category', 'progress').upper()
            time_range = topic.get('time_range', '')
            doc.add_heading(f"{time_range}  {topic.get('topic_title', '')}  [{cat}]", level=2)
            participants = topic.get('participants', [])
            if participants:
                doc.add_paragraph(f"Participants: {', '.join(participants)}")
            doc.add_paragraph(topic.get('summary', ''))
            decisions = topic.get('key_decisions', [])
            if decisions:
                doc.add_heading('Key Decisions', level=3)
                for d in decisions:
                    doc.add_paragraph(d, style='List Bullet')
            actions = topic.get('action_items', [])
            if actions:
                doc.add_heading('Action Items', level=3)
                for ai in actions:
                    priority = ai.get('priority', 'medium').upper()
                    text = (f"[{priority}] {ai.get('action', '')} "
                            f"\u2192 {ai.get('responsible', '?')} by {ai.get('deadline', '?')}")
                    doc.add_paragraph(text, style='List Bullet')
            flags = topic.get('safety_flags', [])
            if flags:
                doc.add_heading('Safety Flags', level=3)
                for sf in flags:
                    risk = sf.get('risk_level', 'medium').upper()
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(f"[{risk}] {sf.get('observation', '')}")
                    if risk == 'HIGH':
                        run.font.color.rgb = RGBColor(192, 57, 43)
                        run.bold = True
                    p.add_run(f"\n  Recommended: {sf.get('recommended_action', '')}")
            photos = topic.get('related_photos', [])
            if photos:
                photo_names = []
                for p in photos:
                    if isinstance(p, str):
                        photo_names.append(p)
                    elif isinstance(p, dict):
                        photo_names.append(p.get('filename', p.get('key', '?')))
                doc.add_paragraph(f"Related photos: {', '.join(photo_names)}")

    for section_key, section_title in [
        ('safety_trends', 'Safety Trends'),
        ('progress_highlights', 'Progress Highlights'),
        ('outstanding_actions', 'Outstanding Actions'),
        ('next_week_priorities', 'Next Week Priorities'),
        ('next_month_priorities', 'Next Month Priorities'),
    ]:
        items = report_data.get(section_key, [])
        if items:
            doc.add_heading(section_title, level=1)
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, str):
                        doc.add_paragraph(item, style='List Bullet')
                    elif isinstance(item, dict):
                        text_parts = [f"{v}" for k, v in item.items() if v]
                        doc.add_paragraph(' \u2014 '.join(text_parts), style='List Bullet')

    quality = report_data.get('quality_summary', '')
    if quality:
        doc.add_heading('Quality Summary', level=1)
        doc.add_paragraph(quality)

    meta = report_data.get('_report_metadata', {})
    if meta:
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(
            f"Generated: {meta.get('generated_at', '?')} | "
            f"Model: {meta.get('model', '?')} | "
            f"Recordings: {meta.get('recordings_processed', '?')} | "
            f"Version: {meta.get('version', '?')}"
        )
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(150, 150, 150)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# Process User Data
# ============================================================

def process_user_data(bucket, user_name, target_date, exclude_keys=None):
    logger.info(f"Processing user: {user_name} for {target_date}")

    transcript_prefix = f"transcripts/{user_name}/{target_date}/"
    transcript_objects = list_s3_objects(bucket, transcript_prefix)
    if not transcript_objects:
        transcript_prefix_all = f"transcripts/{user_name}/"
        transcript_objects = [
            obj for obj in list_s3_objects(bucket, transcript_prefix_all)
            if target_date in obj['key']
        ]

    transcripts = []
    all_text_parts = []
    total_words = 0
    total_duration = 0.0

    for obj in transcript_objects:
        key = obj['key']
        filename = os.path.basename(key)
        if not key.endswith('.json'):
            continue
        if target_date not in key:
            continue
        if exclude_keys and key in exclude_keys:
            logger.info(f'    Skipping (meeting): {os.path.basename(key)}')
            continue

        data = download_json_from_s3(bucket, key)
        parsed = parse_transcript(data)
        if parsed and parsed['full_text']:
            timestamp = extract_timestamp_from_filename(filename)
            device = extract_device_from_filename(filename)
            total_words += parsed['word_count']
            total_duration += parsed.get('duration_seconds', 0.0)
            entry = {
                'key': key, 'filename': filename, 'device': device,
                'user': user_name, 'timestamp': timestamp,
                'time_str': timestamp.strftime('%H:%M:%S') if timestamp else '',
                'text': parsed['full_text'],
                'duration_seconds': parsed.get('duration_seconds', 0.0),
                'vad': extract_vad_info_from_filename(filename),
            }
            transcripts.append(entry)
            time_prefix = f"[{entry['time_str']}] " if entry['time_str'] else ""
            all_text_parts.append(f"{time_prefix}{device}: {parsed['full_text']}")

    transcripts.sort(key=lambda x: x['timestamp'] or datetime.min)

    photo_prefix = f"users/{user_name}/pictures/{target_date}/"
    photo_objects = list_s3_objects(bucket, photo_prefix)
    if not photo_objects:
        photo_prefix_all = f"users/{user_name}/pictures/"
        photo_objects = [
            obj for obj in list_s3_objects(bucket, photo_prefix_all)
            if target_date in obj['key']
        ]

    photos = []
    for obj in photo_objects:
        key = obj['key']
        filename = os.path.basename(key)
        if target_date not in key:
            continue
        timestamp = extract_timestamp_from_filename(filename)
        photos.append({
            'key': key, 'filename': filename, 'timestamp': timestamp,
            'time_str': timestamp.strftime('%H:%M:%S') if timestamp else ''
        })

    logger.info(f"  Found {len(transcripts)} transcripts, {len(photos)} photos, "
                f"{total_words} words, {total_duration:.0f}s audio")

    return {
        'transcripts': transcripts, 'photos': photos,
        'text_parts': all_text_parts, 'total_words': total_words,
        'total_duration': total_duration,
    }


# ============================================================
# Stale Report Detection
# ============================================================

def count_transcripts_for_date(bucket, target_date):
    count = 0
    for obj in list_s3_objects(bucket, 'transcripts/'):
        if obj['key'].endswith('.json') and target_date in obj['key']:
            count += 1
    return count

def get_report_transcript_count(bucket, target_date):
    summary_key = f"{REPORT_PREFIX}{target_date}/summary_report.json"
    summary_data = download_json_from_s3(bucket, summary_key)
    if summary_data:
        meta = summary_data.get('_report_metadata', {})
        return meta.get('recordings_processed', 0)
    total_processed = 0
    found_any = False
    prefix = f"{REPORT_PREFIX}{target_date}/"
    for obj in list_s3_objects(bucket, prefix):
        key = obj['key']
        if not key.endswith('.json'):
            continue
        if '_debug' in key or 'summary_report' in key:
            continue
        data = download_json_from_s3(bucket, key)
        if not data:
            continue
        found_any = True
        meta = data.get('_report_metadata', {})
        total_processed += meta.get('recordings_processed', 0)
    return total_processed if found_any else -1

def check_stale_reports(days_back=7):
    stale_dates = []
    nzdt_now = get_nzdt_now()
    for i in range(2, days_back + 1):
        check_date = (nzdt_now - timedelta(days=i)).strftime('%Y-%m-%d')
        current_transcripts = count_transcripts_for_date(S3_BUCKET, check_date)
        if current_transcripts == 0:
            continue
        report_count = get_report_transcript_count(S3_BUCKET, check_date)
        if report_count == -1:
            logger.info(f"  Backfill {check_date}: NO report, {current_transcripts} transcripts \u2192 GENERATE")
            stale_dates.append(check_date)
        elif current_transcripts > report_count:
            logger.info(f"  Backfill {check_date}: report has {report_count}, "
                        f"now {current_transcripts} transcripts \u2192 STALE, regenerate")
            stale_dates.append(check_date)
        else:
            logger.info(f"  Backfill {check_date}: up to date "
                        f"({current_transcripts} transcripts = {report_count} in report)")
    return stale_dates


# ============================================================
# Report Freshness Check
# ============================================================

def check_report_freshness(bucket, target_date, user_name, current_transcript_count):
    json_key = f"{REPORT_PREFIX}{target_date}/{user_name}/daily_report.json"
    try:
        obj = s3_client.get_object(Bucket=bucket, Key=json_key)
        existing = json.loads(obj['Body'].read().decode('utf-8'))
        meta = existing.get('_report_metadata', {})
        existing_count = meta.get('recordings_processed', -1)
        if existing_count == current_transcript_count:
            return True, existing_count
        else:
            return False, existing_count
    except s3_client.exceptions.NoSuchKey:
        return False, 0
    except Exception as e:
        logger.warning(f"  Freshness check failed for {user_name}: {e}")
        return False, 0

def check_summary_freshness(bucket, target_date, current_transcript_count):
    json_key = f"{REPORT_PREFIX}{target_date}/summary_report.json"
    try:
        obj = s3_client.get_object(Bucket=bucket, Key=json_key)
        existing = json.loads(obj['Body'].read().decode('utf-8'))
        meta = existing.get('_report_metadata', {})
        existing_count = meta.get('recordings_processed', -1)
        if existing_count == current_transcript_count:
            return True, existing_count
        else:
            return False, existing_count
    except s3_client.exceptions.NoSuchKey:
        return False, 0
    except Exception as e:
        logger.warning(f"  Summary freshness check failed: {e}")
        return False, 0


# ============================================================
# Generate Daily Report
# ============================================================

def generate_daily_report(target_date, hidden_topic_ids=None, triggered_by='system',
                          users_filter=None, force=False):
    if users_filter:
        logger.info(f"=== Generating DAILY report for {target_date} \u2014 users: {users_filter} ===")
    else:
        logger.info(f"=== Generating DAILY report for {target_date} ===")

    load_prompt_templates(S3_BUCKET)
    user_mapping = load_user_mapping(S3_BUCKET)
    site_name = os.environ.get('SITE_NAME', 'Construction Site')

    users = set()
    for obj in list_s3_objects(S3_BUCKET, 'transcripts/'):
        parts = obj['key'].split('/')
        if len(parts) >= 2 and parts[1]:
            users.add(parts[1])

    if users_filter:
        users = users & set(users_filter)
        if not users:
            logger.warning(f"No matching users found for filter {users_filter}")

    logger.info(f"Found {len(users)} users: {sorted(users)}")

    user_primary_site, user_all_sites, user_roles, sites_info = get_user_site_mapping(S3_BUCKET)

    all_users_data = {}
    combined_transcripts = []
    combined_photos = []

    # --- Read meeting manifests: skip transcripts consumed by meeting minutes ---
    meeting_consumed_keys = set()
    for user_name in sorted(users):
        manifest_keys = read_meeting_manifest(
            s3_client, S3_BUCKET, REPORT_PREFIX, target_date, user_name)
        if manifest_keys:
            meeting_consumed_keys |= manifest_keys
            logger.info(f'  Meeting manifest for {user_name}: {len(manifest_keys)} transcripts excluded')

    for user_name in sorted(users):
        user_data = process_user_data(S3_BUCKET, user_name, target_date,
                                       exclude_keys=meeting_consumed_keys)
        if user_data['transcripts'] or user_data['photos']:
            all_users_data[user_name] = user_data
            combined_transcripts.extend(user_data['transcripts'])
            combined_photos.extend(user_data['photos'])

    logger.info(f"Users with data for {target_date}: {list(all_users_data.keys())}")

    skipped_users = []
    for user_name, user_data in all_users_data.items():
        if not user_data['transcripts']:
            continue

        n_transcripts = len(user_data['transcripts'])

        if not force:
            is_fresh, existing_count = check_report_freshness(
                S3_BUCKET, target_date, user_name, n_transcripts)
            if is_fresh:
                logger.info(f"  {user_name}: SKIP \u2014 report up-to-date ({existing_count} transcripts)")
                skipped_users.append(user_name)
                continue
            elif existing_count > 0:
                logger.info(f"  {user_name}: STALE \u2014 {existing_count} \u2192 {n_transcripts} transcripts, regenerating")
        else:
            logger.info(f"  {user_name}: FORCE mode \u2014 regenerating regardless")

        correlated = correlate_photos_with_transcripts(user_data['transcripts'], user_data['photos'])
        device = user_data['transcripts'][0].get('device', 'Unknown') if user_data['transcripts'] else 'Unknown'

        user_role = user_roles.get(user_name, '')
        user_site_id = user_primary_site.get(user_name, '')
        user_site_info = sites_info.get(user_site_id, {})
        user_site_name = user_site_info.get('name', site_name)

        prompt = build_daily_prompt(
            correlated, user_name, user_site_name, target_date,
            role=user_role, total_duration=user_data['total_duration'],
            num_photos=len(user_data['photos']), name_mapping=user_mapping,
        )

        max_tokens = min(4096 + n_transcripts * 350, 16000)
        logger.info(f"  {user_name}: {n_transcripts} transcripts \u2192 max_tokens={max_tokens}")

        raw_response, error = call_claude_structured(prompt, max_tokens=max_tokens)

        if error:
            logger.error(f"Claude error for {user_name}: {error}")
            save_debug_record(
                S3_BUCKET, target_date, user_name, prompt, None, None,
                parse_success=False,
                input_stats={
                    'transcripts_count': n_transcripts,
                    'total_words': user_data['total_words'],
                    'total_duration': user_data['total_duration'],
                    'photos_count': len(user_data['photos']),
                    'error': error,
                }
            )
            continue

        claude_output = extract_json_from_response(raw_response)
        parse_success = claude_output is not None

        if not claude_output:
            logger.error(f"Failed to parse Claude JSON for {user_name}, saving raw text")
            claude_output = {
                'executive_summary': raw_response[:500] if raw_response else 'Error generating report',
                'topics': [], 'safety_observations': [],
            }

        save_debug_record(
            S3_BUCKET, target_date, user_name, prompt, raw_response, claude_output,
            parse_success=parse_success,
            input_stats={
                'transcripts_count': n_transcripts,
                'total_words': user_data['total_words'],
                'total_duration': user_data['total_duration'],
                'photos_count': len(user_data['photos']),
            }
        )

        topics = claude_output.get('topics', [])
        if hidden_topic_ids:
            topics = [t for t in topics if t.get('topic_id') not in hidden_topic_ids]
            claude_output['topics'] = topics

        dur_total = user_data['total_duration']
        dur_min = int(dur_total // 60)
        dur_sec = int(dur_total % 60)

        recording_durations = []
        for t in user_data['transcripts']:
            t_dur = t.get('duration_seconds', 0)
            t_min = int(t_dur // 60)
            t_sec = int(t_dur % 60)
            recording_durations.append({
                'filename': t.get('filename', ''),
                'time': t.get('time_str', ''),
                'duration_seconds': round(t_dur, 1),
                'duration_display': f"{t_min}m {t_sec}s",
            })

        now_iso = datetime.utcnow().isoformat() + 'Z'
        report = {
            'report_date': target_date,
            'report_type': 'daily',
            'user_name': user_name,
            'device': device,
            'site': user_site_name,
            'recording_session': {
                'date': target_date, 'site': user_site_name,
                'worker': user_name, 'role': user_role or 'worker',
                'recordings': n_transcripts,
                'total_duration_seconds': round(dur_total, 1),
                'total_duration_display': f"{dur_min}m {dur_sec}s",
                'total_words': user_data['total_words'],
                'photos': len(user_data['photos']),
                'per_recording': recording_durations,
            },
            'executive_summary': claude_output.get('executive_summary', ''),
            'quality_and_compliance': claude_output.get('quality_and_compliance', []),
            'safety_observations': claude_output.get('safety_observations', []),
            'critical_dates_and_deadlines': claude_output.get('critical_dates_and_deadlines', []),
            'topics': topics,
            '_report_metadata': {
                'version': 'v3.5',
                'generated_at': now_iso,
                'generated_by': triggered_by,
                'recordings_processed': n_transcripts,
                'recordings_skipped': 0,
                'total_words': user_data['total_words'],
                'total_duration_seconds': round(dur_total, 1),
                'model': CLAUDE_MODEL,
                'hidden_topic_ids': hidden_topic_ids or [],
                'parse_success': parse_success,
            }
        }

        json_key = f"{REPORT_PREFIX}{target_date}/{user_name}/daily_report.json"
        s3_client.put_object(
            Bucket=S3_BUCKET, Key=json_key,
            Body=json.dumps(report, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json'
        )
        logger.info(f"Saved: {json_key}")

        try:
            word_buffer = generate_word_document(report, f"Daily Report \u2014 {user_name} \u2014 {target_date}")
            if word_buffer:
                word_key = f"{REPORT_PREFIX}{target_date}/{user_name}/daily_report.docx"
                s3_client.put_object(
                    Bucket=S3_BUCKET, Key=word_key,
                    Body=word_buffer.getvalue(),
                    ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                logger.info(f"Saved: {word_key}")
        except Exception as e:
            logger.error(f"Word generation failed for {user_name}: {e}")

        site_id = user_site_id or site_name.lower().replace(' ', '-')
        write_items_to_dynamodb(site_id, target_date, topics, user_name, device)
        write_report_to_dynamodb(site_id, target_date, 'daily', json_key, user_name,
            visible_count=len(topics),
            hidden_count=len(hidden_topic_ids) if hidden_topic_ids else 0)
        write_audit_entry(site_id, target_date, 'generate_report',
            f"Daily \u2014 {len(topics)} topics, {len(hidden_topic_ids or [])} hidden, "
            f"{n_transcripts} recordings, triggered_by={triggered_by}",
            user=triggered_by)

    # --- Combined summary report ---
    if combined_transcripts and not users_filter:
        if len(skipped_users) == len(all_users_data) and not force:
            logger.info(f"  Combined summary: SKIP \u2014 all {len(skipped_users)} users up-to-date")
        else:
            total_transcript_count = len(combined_transcripts)
            skip_summary = False
            if not force:
                is_fresh, existing_count = check_summary_freshness(
                    S3_BUCKET, target_date, total_transcript_count)
                if is_fresh:
                    logger.info(f"  Combined summary: SKIP \u2014 up-to-date ({existing_count} transcripts)")
                    skip_summary = True

            if not skip_summary:
                combined_transcripts.sort(key=lambda x: x['timestamp'] or datetime.min)
                correlated_all = correlate_photos_with_transcripts(combined_transcripts, combined_photos)
                combined_duration = sum(d['total_duration'] for d in all_users_data.values())
                prompt = build_daily_prompt(
                    correlated_all, "All Workers", site_name, target_date,
                    total_duration=combined_duration,
                    num_photos=len(combined_photos), name_mapping=user_mapping,
                )
                max_tokens = min(4096 + len(combined_transcripts) * 350, 16000)
                raw_response, error = call_claude_structured(prompt, max_tokens=max_tokens)

                if not error:
                    claude_output = extract_json_from_response(raw_response)
                    save_debug_record(
                        S3_BUCKET, target_date, '_summary', prompt, raw_response, claude_output,
                        parse_success=claude_output is not None,
                        input_stats={
                            'transcripts_count': len(combined_transcripts),
                            'total_words': sum(d['total_words'] for d in all_users_data.values()),
                            'total_duration': combined_duration,
                            'photos_count': len(combined_photos),
                            'users': list(all_users_data.keys()),
                        }
                    )

                    if claude_output:
                        comb_dur_min = int(combined_duration // 60)
                        comb_dur_sec = int(combined_duration % 60)
                        combined_report = {
                            'report_date': target_date, 'report_type': 'daily',
                            'user_name': None, 'site': site_name,
                            'users': list(all_users_data.keys()),
                            'recording_session': {
                                'date': target_date, 'site': site_name,
                                'workers': list(all_users_data.keys()),
                                'recordings': len(combined_transcripts),
                                'total_duration_seconds': round(combined_duration, 1),
                                'total_duration_display': f"{comb_dur_min}m {comb_dur_sec}s",
                                'total_words': sum(d['total_words'] for d in all_users_data.values()),
                                'photos': len(combined_photos),
                            },
                            'executive_summary': claude_output.get('executive_summary', ''),
                            'quality_and_compliance': claude_output.get('quality_and_compliance', []),
                            'safety_observations': claude_output.get('safety_observations', []),
                            'critical_dates_and_deadlines': claude_output.get('critical_dates_and_deadlines', []),
                            'topics': claude_output.get('topics', []),
                            '_report_metadata': {
                                'version': 'v3.5',
                                'generated_at': datetime.utcnow().isoformat() + 'Z',
                                'generated_by': triggered_by,
                                'recordings_processed': len(combined_transcripts),
                                'total_words': sum(d['total_words'] for d in all_users_data.values()),
                                'total_duration_seconds': round(combined_duration, 1),
                                'model': CLAUDE_MODEL,
                            }
                        }

                        json_key = f"{REPORT_PREFIX}{target_date}/summary_report.json"
                        s3_client.put_object(
                            Bucket=S3_BUCKET, Key=json_key,
                            Body=json.dumps(combined_report, ensure_ascii=False, indent=2, default=str),
                            ContentType='application/json'
                        )
                        logger.info(f"Saved: {json_key}")

                        try:
                            word_buffer = generate_word_document(combined_report, f"Daily Summary \u2014 {target_date}")
                            if word_buffer:
                                word_key = f"{REPORT_PREFIX}{target_date}/summary_report.docx"
                                s3_client.put_object(
                                    Bucket=S3_BUCKET, Key=word_key,
                                    Body=word_buffer.getvalue(),
                                    ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                                )
                                logger.info(f"Saved: {word_key}")
                        except Exception as e:
                            logger.error(f"Combined Word failed: {e}")

    return {
        'report_type': 'daily', 'date': target_date,
        'users_processed': list(all_users_data.keys()),
        'users_skipped': skipped_users,
        'total_transcripts': len(combined_transcripts),
        'total_photos': len(combined_photos),
        'total_duration_seconds': round(sum(d['total_duration'] for d in all_users_data.values()), 1),
        'force': force,
    }


# ============================================================
# Generate Weekly / Monthly Report
# ============================================================

def generate_periodic_report(report_type, start_date, end_date):
    logger.info(f"=== Generating {report_type.upper()} report: {start_date} to {end_date} ===")
    load_prompt_templates(S3_BUCKET)
    site_name = os.environ.get('SITE_NAME', 'Construction Site')
    site_id_default = site_name.lower().replace(' ', '-')
    user_primary_site, user_all_sites, user_roles, sites_info = get_user_site_mapping(S3_BUCKET)
    logger.info(f"  User-site mapping: {len(user_primary_site)} users across "
                f"{len(set(user_primary_site.values()))} sites")

    all_daily_reports = []
    reports_by_user = {}

    for date_str in dates_in_range(start_date, end_date):
        prefix = f"{REPORT_PREFIX}{date_str}/"
        found_per_user = False
        for obj in list_s3_objects(S3_BUCKET, prefix):
            key = obj['key']
            if key.endswith('/daily_report.json') and '/sites/' not in key:
                data = download_json_from_s3(S3_BUCKET, key)
                if data and data.get('user_name'):
                    all_daily_reports.append(data)
                    user = data['user_name']
                    reports_by_user.setdefault(user, []).append(data)
                    found_per_user = True
            elif '/by_user/' in key and key.endswith('.json') and '_debug' not in key:
                data = download_json_from_s3(S3_BUCKET, key)
                if data:
                    user = data.get('user_name', data.get('user', 'Unknown'))
                    all_daily_reports.append(data)
                    reports_by_user.setdefault(user, []).append(data)
                    found_per_user = True
        if not found_per_user:
            summary_key = f"{REPORT_PREFIX}{date_str}/summary_report.json"
            data = download_json_from_s3(S3_BUCKET, summary_key)
            if data:
                all_daily_reports.append(data)
                reports_by_user.setdefault('_summary', []).append(data)

    logger.info(f"  Collected {len(all_daily_reports)} daily reports, "
                f"{len(reports_by_user)} unique users: {sorted(reports_by_user.keys())}")

    if not all_daily_reports:
        logger.warning("  No daily reports found for period")
        return {'report_type': report_type, 'status': 'no_data'}

    # Fetch and apply user corrections (QA/QC Layer 2)
    corrections_applied = 0
    try:
        corrections_map = fetch_corrections_for_range(start_date, end_date)
        if corrections_map:
            total_corrections = sum(len(v) for v in corrections_map.values())
            logger.info(f"  Fetched {total_corrections} corrections across {len(corrections_map)} dates")
            corrections_applied = apply_corrections_to_reports(all_daily_reports, corrections_map)
            logger.info(f"  Applied {corrections_applied} corrections to daily reports")
    except Exception as e:
        logger.warning(f"  Corrections fetch/apply failed (continuing without): {e}")

    weekly_reports = []
    if report_type == 'monthly':
        for obj in list_s3_objects(S3_BUCKET, REPORT_PREFIX):
            if 'weekly_report.json' in obj['key']:
                wr = download_json_from_s3(S3_BUCKET, obj['key'])
                if wr:
                    weekly_reports.append(wr)

    now_iso = datetime.utcnow().isoformat() + 'Z'
    per_user_results = {}

    for user_name, user_reports in reports_by_user.items():
        if user_name == '_summary' or not user_reports:
            continue
        logger.info(f"  Generating {report_type} report for user: {user_name} ({len(user_reports)} daily reports)")
        user_site_id = user_primary_site.get(user_name, '')
        user_site_info = sites_info.get(user_site_id, {})
        user_site_name = user_site_info.get('name', site_name)
        user_role = user_roles.get(user_name, '')

        if report_type == 'weekly':
            prompt = build_weekly_prompt(user_reports, user_site_name, start_date, end_date,
                scope_label=user_name, scope_type='user', user_role=user_role,
                corrections_applied=corrections_applied)
        else:
            prompt = build_monthly_prompt(user_reports, [], user_site_name, start_date, end_date,
                corrections_applied=corrections_applied)

        raw_response, error = call_claude_structured(prompt, max_tokens=6000)
        if error:
            per_user_results[user_name] = {'status': 'error', 'error': error}
            continue
        claude_output = extract_json_from_response(raw_response)
        if not claude_output:
            per_user_results[user_name] = {'status': 'parse_error'}
            continue

        user_report = {
            'report_date': end_date, 'report_type': report_type,
            'period': {'start': start_date, 'end': end_date},
            'user_name': user_name, 'role': user_role,
            'site': user_site_name, 'site_id': user_site_id,
            **claude_output,
            '_report_metadata': {
                'version': 'v3.5', 'generated_at': now_iso, 'generated_by': 'system',
                'scope': 'user', 'daily_reports_used': len(user_reports), 'model': CLAUDE_MODEL,
                'corrections_applied': corrections_applied,
            }
        }

        safe_user = user_name.replace(' ', '_')
        json_key = f"{REPORT_PREFIX}{end_date}/{safe_user}/{report_type}_report.json"
        s3_client.put_object(Bucket=S3_BUCKET, Key=json_key,
            Body=json.dumps(user_report, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json')
        logger.info(f"  Saved: {json_key}")

        try:
            title = f"{'Weekly' if report_type == 'weekly' else 'Monthly'} Report \u2014 {user_name} \u2014 {start_date} to {end_date}"
            word_buffer = generate_word_document(user_report, title)
            if word_buffer:
                word_key = f"{REPORT_PREFIX}{end_date}/{safe_user}/{report_type}_report.docx"
                s3_client.put_object(Bucket=S3_BUCKET, Key=word_key,
                    Body=word_buffer.getvalue(),
                    ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                logger.info(f"  Saved: {word_key}")
        except Exception as e:
            logger.error(f"  Word failed for {user_name}: {e}")

        per_user_results[user_name] = {'status': 'success', 'daily_reports': len(user_reports)}

    # Per-site reports
    reports_by_site = {}
    for user_name, user_reports in reports_by_user.items():
        if user_name == '_summary':
            continue
        user_site_list = user_all_sites.get(user_name, [])
        if not user_site_list:
            user_site_list = [site_id_default]
        for sid in user_site_list:
            reports_by_site.setdefault(sid, []).extend(user_reports)

    per_site_results = {}
    for site_id, site_reports in reports_by_site.items():
        site_info = sites_info.get(site_id, {})
        site_display_name = site_info.get('name', site_name)
        site_location = site_info.get('location', '')
        site_client = site_info.get('client', '')
        site_label_for_prompt = site_display_name
        if site_location or site_client:
            detail_parts = [p for p in [site_location, site_client] if p]
            site_label_for_prompt += f" ({', '.join(detail_parts)})"

        logger.info(f"  Generating {report_type} report for site: {site_display_name} ({len(site_reports)} daily reports)")

        if report_type == 'weekly':
            prompt = build_weekly_prompt(site_reports, site_display_name, start_date, end_date,
                scope_label=site_label_for_prompt, scope_type='site',
                corrections_applied=corrections_applied)
        else:
            prompt = build_monthly_prompt(site_reports, weekly_reports, site_display_name, start_date, end_date,
                corrections_applied=corrections_applied)

        raw_response, error = call_claude_structured(prompt, max_tokens=6000)
        if error:
            per_site_results[site_id] = {'status': 'error', 'error': error}
            continue
        claude_output = extract_json_from_response(raw_response)
        if not claude_output:
            per_site_results[site_id] = {'status': 'parse_error'}
            continue

        site_report = {
            'report_date': end_date, 'report_type': report_type,
            'period': {'start': start_date, 'end': end_date},
            'site': site_display_name, 'site_id': site_id,
            'site_location': site_info.get('location', ''),
            'site_client': site_info.get('client', ''),
            'users': list(set(r.get('user_name', '?') for r in site_reports)),
            **claude_output,
            '_report_metadata': {
                'version': 'v3.5', 'generated_at': now_iso, 'generated_by': 'system',
                'scope': 'site', 'daily_reports_used': len(site_reports), 'model': CLAUDE_MODEL,
                'corrections_applied': corrections_applied,
            }
        }

        json_key = f"{REPORT_PREFIX}{end_date}/sites/{site_id}/{report_type}_report.json"
        s3_client.put_object(Bucket=S3_BUCKET, Key=json_key,
            Body=json.dumps(site_report, ensure_ascii=False, indent=2, default=str),
            ContentType='application/json')
        logger.info(f"  Saved: {json_key}")

        try:
            title = f"{'Weekly' if report_type == 'weekly' else 'Monthly'} Report \u2014 {site_display_name} \u2014 {start_date} to {end_date}"
            word_buffer = generate_word_document(site_report, title)
            if word_buffer:
                word_key = f"{REPORT_PREFIX}{end_date}/sites/{site_id}/{report_type}_report.docx"
                s3_client.put_object(Bucket=S3_BUCKET, Key=word_key,
                    Body=word_buffer.getvalue(),
                    ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                logger.info(f"  Saved: {word_key}")
        except Exception as e:
            logger.error(f"  Site Word failed for {site_id}: {e}")

        write_report_to_dynamodb(site_id, end_date, report_type, json_key, None,
            visible_count=len(site_reports), hidden_count=0)
        write_audit_entry(site_id, end_date, 'generate_report',
            f"{report_type.capitalize()} site report \u2014 {len(site_reports)} daily reports")
        per_site_results[site_id] = {'status': 'success', 'daily_reports': len(site_reports)}

    # Combined summary
    if len(reports_by_site) > 1 or True:
        logger.info(f"  Generating combined {report_type} summary ({len(all_daily_reports)} total daily reports)")
        if report_type == 'weekly':
            prompt = build_weekly_prompt(all_daily_reports, site_name, start_date, end_date,
                scope_label=None, scope_type='all',
                corrections_applied=corrections_applied)
        else:
            prompt = build_monthly_prompt(all_daily_reports, weekly_reports, site_name, start_date, end_date,
                corrections_applied=corrections_applied)

        raw_response, error = call_claude_structured(prompt, max_tokens=6000)
        if not error:
            claude_output = extract_json_from_response(raw_response)
            if claude_output:
                combined_report = {
                    'report_date': end_date, 'report_type': report_type,
                    'period': {'start': start_date, 'end': end_date},
                    'site': site_name,
                    'users': list(set(r.get('user_name', '?') for r in all_daily_reports if r.get('user_name'))),
                    **claude_output,
                    '_report_metadata': {
                        'version': 'v3.5', 'generated_at': now_iso, 'generated_by': 'system',
                        'scope': 'all', 'daily_reports_used': len(all_daily_reports), 'model': CLAUDE_MODEL,
                        'corrections_applied': corrections_applied,
                    }
                }
                json_key = f"{REPORT_PREFIX}{end_date}/{report_type}_report.json"
                s3_client.put_object(Bucket=S3_BUCKET, Key=json_key,
                    Body=json.dumps(combined_report, ensure_ascii=False, indent=2, default=str),
                    ContentType='application/json')
                logger.info(f"  Saved: {json_key}")
                try:
                    title = f"{'Weekly' if report_type == 'weekly' else 'Monthly'} Report \u2014 {start_date} to {end_date}"
                    word_buffer = generate_word_document(combined_report, title)
                    if word_buffer:
                        word_key = f"{REPORT_PREFIX}{end_date}/{report_type}_report.docx"
                        s3_client.put_object(Bucket=S3_BUCKET, Key=word_key,
                            Body=word_buffer.getvalue(),
                            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        logger.info(f"  Saved: {word_key}")
                except Exception as e:
                    logger.error(f"Combined Word failed: {e}")

    return {
        'report_type': report_type,
        'period': {'start': start_date, 'end': end_date},
        'daily_reports_used': len(all_daily_reports),
        'per_user': per_user_results,
        'per_site': per_site_results,
        'status': 'success',
    }


# ============================================================
# MAIN HANDLER
# ============================================================

def lambda_handler(event, context):
    """
    Event payload options:
      {"report_type": "daily"}
      {"report_type": "daily", "date": "2026-02-19"}
      {"report_type": "daily", "date": "2026-02-23", "user": "David_Barillaro"}
      {"report_type": "daily", "date": "...", "user": ["MPI1", "MPI3"]}
      {"report_type": "daily", "date": "...", "hidden_topic_ids": [3], "triggered_by": "user@email"}
      {"report_type": "daily", "skip_backfill": true}
      {"report_type": "daily", "force": true}
      {"report_type": "weekly"}
      {"report_type": "weekly", "start_date": "...", "end_date": "..."}
      {"report_type": "monthly"}
      {"report_type": "monthly", "start_date": "...", "end_date": "..."}
    """
    logger.info("=" * 60)
    logger.info("Report Generator v3.5 - Starting")
    logger.info(f"Event: {json.dumps(event, default=str)}")
    logger.info(f"Word generation: {'enabled' if DOCX_AVAILABLE else 'DISABLED (no python-docx layer)'}")
    logger.info(f"DynamoDB: {'enabled' if ENABLE_DYNAMODB else 'DISABLED (set ENABLE_DYNAMODB=true to enable)'}")
    logger.info("=" * 60)

    if not S3_BUCKET:
        return {'statusCode': 400, 'body': 'Missing S3_BUCKET'}

    global _user_mapping_cache
    _user_mapping_cache = None
    global _prompt_templates_cache
    _prompt_templates_cache = None

    report_type = event.get('report_type', 'daily')

    if report_type == 'daily':
        target_date = event.get('date', get_yesterday_date())
        hidden_ids = event.get('hidden_topic_ids', None)
        triggered_by = event.get('triggered_by', 'system')
        force = event.get('force', False)

        users_filter = event.get('user', None)
        if isinstance(users_filter, str):
            users_filter = [users_filter]

        result = generate_daily_report(target_date, hidden_ids, triggered_by,
                                       users_filter=users_filter, force=force)

        if not event.get('skip_backfill'):
            logger.info(f"Checking past {BACKFILL_DAYS} days for stale reports...")
            stale_dates = check_stale_reports(days_back=BACKFILL_DAYS)
            if stale_dates:
                logger.info(f"Found {len(stale_dates)} stale reports: {stale_dates}")
                backfill_results = []
                for stale_date in stale_dates:
                    try:
                        logger.info(f"Regenerating report for {stale_date}...")
                        r = generate_daily_report(stale_date, triggered_by='backfill')
                        backfill_results.append({
                            'date': stale_date, 'status': 'success',
                            'users': r.get('users_processed', []),
                            'transcripts': r.get('total_transcripts', 0),
                        })
                    except Exception as e:
                        logger.error(f"Backfill failed for {stale_date}: {e}")
                        backfill_results.append({'date': stale_date, 'status': 'error', 'error': str(e)})
                result['backfill'] = backfill_results
            else:
                logger.info("All past reports are up to date.")
                result['backfill'] = []

    elif report_type == 'weekly':
        start = event.get('start_date')
        end = event.get('end_date')
        if not start or not end:
            start, end = get_week_range()
        result = generate_periodic_report('weekly', start, end)

    elif report_type == 'monthly':
        start = event.get('start_date')
        end = event.get('end_date')
        if not start or not end:
            start, end = get_month_range()
        result = generate_periodic_report('monthly', start, end)

    else:
        return {'statusCode': 400, 'body': f'Unknown report_type: {report_type}'}

    logger.info("=" * 60)
    logger.info(f"Report Generation Complete: {json.dumps(result, default=str)}")
    logger.info("=" * 60)

    return {'statusCode': 200, 'body': json.dumps(result, default=str)}